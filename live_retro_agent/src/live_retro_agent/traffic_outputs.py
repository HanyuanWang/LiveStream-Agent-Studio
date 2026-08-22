from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from .config import Settings


COLORS = {
    "dark": "0F5B3C", "green": "177A50", "mint": "E7F3EC", "orange": "F47A4A",
    "pale_orange": "FFF0E8", "gold": "8A5A10", "pale_gold": "FFF6D9", "red": "B42318",
    "pale_red": "FDECEC", "blue": "2E6F9E", "pale_blue": "EAF3F8", "ink": "17231D",
    "gray": "667085", "pale": "F7F8F6", "grid": "D9E1DC", "white": "FFFFFF",
}


def build_xlsx(path: Path, payload: dict[str, Any]) -> None:
    workbook = Workbook()
    workbook.remove(workbook.active)
    thin = Side(style="thin", color=COLORS["grid"])
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for definition in payload.get("sheets", []):
        sheet = workbook.create_sheet(str(definition.get("name") or "Sheet")[:31])
        sheet.sheet_view.showGridLines = False
        freeze_rows = int(definition.get("freezeRows") or 0)
        if freeze_rows:
            sheet.freeze_panes = f"A{freeze_rows + 1}"
        for index, width in enumerate(definition.get("widths") or [], 1):
            sheet.column_dimensions[get_column_letter(index)].width = float(width)
        row = 1
        max_columns = max([2] + [len(block.get("headers") or []) for block in definition.get("blocks", [])])
        for block in definition.get("blocks", []):
            block_type = block.get("type")
            if block_type in {"title", "section", "note"}:
                sheet.merge_cells(start_row=row, start_column=1, end_row=row, end_column=max_columns)
                cell = sheet.cell(row, 1, text(block.get("text")))
                if block_type == "title":
                    cell.fill = PatternFill("solid", fgColor=COLORS["dark"])
                    cell.font = Font(name="Microsoft YaHei", bold=True, color=COLORS["white"], size=18)
                    sheet.row_dimensions[row].height = 40
                    row += 2
                elif block_type == "section":
                    tone = block.get("tone")
                    fill = COLORS["pale_gold"] if tone == "warning" else COLORS["pale_red"] if tone == "danger" else COLORS["pale_blue"] if tone == "blue" else COLORS["mint"]
                    color = COLORS["gold"] if tone == "warning" else COLORS["red"] if tone == "danger" else COLORS["blue"] if tone == "blue" else COLORS["dark"]
                    cell.fill = PatternFill("solid", fgColor=fill)
                    cell.font = Font(name="Microsoft YaHei", bold=True, color=color, size=12)
                    sheet.row_dimensions[row].height = 29
                    row += 1
                else:
                    cell.fill = PatternFill("solid", fgColor=COLORS["pale"])
                    cell.font = Font(name="Microsoft YaHei", color=COLORS["ink"], size=10)
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    sheet.row_dimensions[row].height = max(38, int(block.get("lines") or 3) * 15)
                    row += 1
                continue
            if block_type != "table":
                continue
            headers = list(block.get("headers") or [])
            table_rows = list(block.get("rows") or [])
            for column, header in enumerate(headers, 1):
                cell = sheet.cell(row, column, text(header))
                cell.fill = PatternFill("solid", fgColor=COLORS["dark"])
                cell.font = Font(name="Microsoft YaHei", bold=True, color=COLORS["white"], size=9)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = border
            sheet.row_dimensions[row].height = 28
            row += 1
            for values in table_rows:
                for column, value in enumerate(values, 1):
                    cell = sheet.cell(row, column, value if value is not None else "缺失")
                    cell.font = Font(name="Microsoft YaHei", color=COLORS["ink"], size=9)
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    cell.border = border
                sheet.row_dimensions[row].height = float(block.get("rowHeight") or 38)
                row += 1
            row += 1
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def items(value: Any) -> list[Any]:
    return value if isinstance(value, list) else ([] if value in (None, "") else [value])


def text(value: Any) -> str:
    if value in (None, ""):
        return "未识别"
    return json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value)


def rows_for(value: Any, fields: list[str]) -> list[list[str]]:
    result = []
    for item in items(value):
        if isinstance(item, dict):
            result.append([text(item.get(field)) for field in fields])
        else:
            result.append([text(item)] + ["未识别"] * (len(fields) - 1))
    return result


def build_payload(sentences: list[dict[str, Any]], timeline: dict[str, Any], review: dict[str, Any]) -> dict[str, Any]:
    conclusions = review["conclusions"]
    cycle = conclusions.get("cycle_recommendation") if isinstance(conclusions.get("cycle_recommendation"), dict) else {}
    minutes = timeline.get("minutes", [])
    traffic_overview = [[
        len(minutes),
        sum(float(m.get("viewers", 0) or 0) for m in minutes),
        sum(float(m.get("leavers", 0) or 0) for m in minutes),
        max([float(m.get("online", 0) or 0) for m in minutes] or [0]),
        round(sum(float(m.get("watch_seconds", 0) or 0) for m in minutes) / max(len(minutes), 1), 1),
        sum(float(m.get("interaction", 0) or 0) for m in minutes),
        sum(float(m.get("followers", 0) or 0) for m in minutes),
        sum(float(m.get("exposure", 0) or 0) for m in minutes),
        sum(float(m.get("product_clicks", 0) or 0) for m in minutes),
    ]]
    traffic_tasks = [
        ["拉进入后的第一停留", "具体身份、单一场景和可回答问题"],
        ["提升停留", "可观察细节加一条清晰证明链"],
        ["提升互动与关注", "给出低门槛动作，并明确用户能得到什么"],
        ["提升商品点击", "证明讲清后再给明确动作，避免只重复库存口令"],
    ]
    verification_metrics = [[i + 1, value] for i, value in enumerate([
        "每轮的进入、离开与净流入", "循环开始、结束和峰值实时在线", "平均观看时长",
        "评论、点赞和新增关注", "商品曝光、商品点击及点击/曝光",
        "本轮使用的主场景、个人证明、核心证明和行动表达",
    ])]
    summary = [
        {"type": "title", "text": "直播流量与话术关系分析"},
        {"type": "section", "text": "一、核心结论"},
        {"type": "note", "text": text(conclusions.get("headline")), "lines": 5},
        {"type": "section", "text": "二、整场流量概览"},
        {"type": "table", "headers": ["数据分钟数", "进入合计", "离开合计", "最高在线", "平均观看秒", "互动合计", "新增关注", "商品曝光", "商品点击"], "rows": traffic_overview, "rowHeight": 54},
        {"type": "section", "text": "三、表现较好的流量波次"},
        {"type": "table", "headers": ["时间段", "流量与话术分析", "复用/执行动作"], "rows": rows_for(conclusions.get("strong_period_analysis"), ["time", "analysis", "action"]), "rowHeight": 110},
        {"type": "section", "text": "四、表现较弱的流量波次", "tone": "warning"},
        {"type": "table", "headers": ["时间段", "问题判断", "修正动作"], "rows": rows_for(conclusions.get("weak_period_analysis"), ["time", "analysis", "action"]), "rowHeight": 110},
        {"type": "section", "text": "五、话术分别应该承担什么流量任务"},
        {"type": "table", "headers": ["话术任务", "表达要求"], "rows": traffic_tasks, "rowHeight": 62},
        {"type": "section", "text": "六、建议的话术循环"},
        {"type": "table", "headers": ["建议时长", "判断依据", "环节安排"], "rows": [[text(cycle.get("range")), text(cycle.get("reason")), text(cycle.get("allocation"))]], "rowHeight": 90},
        {"type": "section", "text": "七、下一场可以直接执行的改进"},
        {"type": "table", "headers": ["序号", "执行动作"], "rows": [[i + 1, text(x)] for i, x in enumerate(items(conclusions.get("host_actions")))], "rowHeight": 50},
        {"type": "section", "text": "八、下场验证指标"},
        {"type": "table", "headers": ["序号", "验证指标"], "rows": verification_metrics, "rowHeight": 46},
    ]

    block_rows = [[b.get("video_time"), b.get("entrants"), b.get("leavers"), b.get("net_flow"), b.get("average_online"), b.get("peak_online"), b.get("watch_seconds"), b.get("interaction"), b.get("likes"), b.get("followers"), b.get("exposure"), b.get("clicks"), b.get("click_rate"), b.get("excerpt")] for b in review["blocks"]]
    minute_rows = [[m.get("video_time"), m.get("original_time"), m.get("viewers", 0), m.get("leavers", 0), m.get("online", 0), m.get("watch_seconds", 0), m.get("interaction", 0), m.get("likes", 0), m.get("followers", 0), m.get("exposure", 0), m.get("product_clicks", 0)] for m in timeline["minutes"]]
    sentence_rows = [[s.get("sentence_id"), s.get("video_start"), s.get("video_end"), s.get("event"), s.get("phase"), s.get("text"), s.get("traffic_viewers", 0), s.get("traffic_leavers", 0), s.get("traffic_online", 0), s.get("traffic_watch_seconds", 0), s.get("traffic_interaction", 0), s.get("traffic_likes", 0), s.get("traffic_followers", 0), s.get("traffic_exposure", 0), s.get("traffic_clicks", 0)] for s in sentences]
    raw_rows = [[s.get("source_sheet"), s.get("source_row"), s.get("source_timestamp"), s.get("event"), s.get("text")] for s in sentences]

    return {"sheets": [
        {"name": "结论", "freezeRows": 2, "widths": [18, 82, 65, 70, 18], "blocks": summary},
        {"name": "流量波次", "freezeRows": 4, "widths": [22, 13, 13, 13, 15, 15, 15, 13, 13, 13, 15, 15, 15, 90], "blocks": [{"type": "title", "text": "流量波次与同期话术"}, {"type": "table", "headers": ["时间", "进入", "离开", "净流入", "平均在线", "最高在线", "平均观看秒", "评论/互动", "点赞", "新增关注", "商品曝光", "商品点击", "点击/曝光", "同期话术"], "rows": block_rows, "rowHeight": 92}]},
        {"name": "逐句话术×流量", "freezeRows": 4, "widths": [9, 14, 14, 24, 14, 90, 13, 13, 15, 15, 14, 12, 14, 14, 14], "blocks": [{"type": "title", "text": "逐句话术与同期流量"}, {"type": "table", "headers": ["句号", "开始", "结束", "事件", "话术环节", "主播原话", "进入", "离开", "在线峰值", "平均观看秒", "互动", "点赞", "关注", "商品曝光", "商品点击"], "rows": sentence_rows, "rowHeight": 72}]},
        {"name": "分钟流量", "freezeRows": 4, "widths": [15, 22, 13, 13, 15, 15, 13, 13, 13, 15, 15], "blocks": [{"type": "title", "text": "巨量百应分钟流量数据"}, {"type": "table", "headers": ["视频时间", "原始时间", "进入", "离开", "实时在线", "人均观看秒", "评论/互动", "点赞", "新增关注", "商品曝光", "商品点击"], "rows": minute_rows, "rowHeight": 38}]},
        {"name": "原始逐字稿", "freezeRows": 4, "widths": [20, 10, 24, 28, 100], "blocks": [{"type": "title", "text": "原始逐字稿来源追溯"}, {"type": "table", "headers": ["来源Sheet", "源行", "源时间戳", "事件", "逐字稿"], "rows": raw_rows, "rowHeight": 74}]},
    ]}


def add_analysis_items(document: Document, heading: str, value: Any, labels: list[tuple[str, str]]) -> None:
    document.add_heading(heading, level=1)
    for index, item in enumerate(items(value), 1):
        if not isinstance(item, dict):
            document.add_paragraph(text(item))
            continue
        document.add_heading(f"{index}. {text(item.get(labels[0][0]))}", level=2)
        for key, label in labels[1:]:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}：").bold = True
            paragraph.add_run(text(item.get(key)))


def build_docx(path: Path, review: dict[str, Any]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8); section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    for name in ("Normal", "Title", "Heading 1", "Heading 2"):
        style = document.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    document.styles["Normal"].font.size = Pt(10.5)
    document.styles["Title"].font.size = Pt(25); document.styles["Title"].font.bold = True
    document.styles["Heading 1"].font.size = Pt(15); document.styles["Heading 1"].font.color.rgb = RGBColor(15, 91, 60)
    title = document.add_paragraph(style="Title"); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.add_run("直播流量与话术关系分析")
    c = review["conclusions"]
    document.add_heading("一、核心结论", level=1)
    document.add_paragraph(text(c.get("headline")))
    document.add_heading("二、整场流量概览", level=1)
    minutes = review.get("timeline_minutes", [])
    if minutes:
        overview = document.add_table(rows=2, cols=5)
        overview.style = "Table Grid"
        headers = ["数据分钟数", "进入合计", "离开合计", "最高在线", "平均观看秒"]
        values = [
            len(minutes),
            round(sum(float(m.get("viewers", 0) or 0) for m in minutes)),
            round(sum(float(m.get("leavers", 0) or 0) for m in minutes)),
            round(max([float(m.get("online", 0) or 0) for m in minutes] or [0])),
            round(sum(float(m.get("watch_seconds", 0) or 0) for m in minutes) / max(len(minutes), 1), 1),
        ]
        for index, value in enumerate(headers): overview.cell(0, index).text = value
        for index, value in enumerate(values): overview.cell(1, index).text = str(value)
    else:
        document.add_paragraph("整场流量明细见配套 Excel 的“分钟流量”和“流量波次”工作表。")
    add_analysis_items(document, "三、表现较好的流量波次", c.get("strong_period_analysis"), [("time", "时间"), ("analysis", "分析"), ("action", "复用/执行动作")])
    add_analysis_items(document, "四、表现较弱的流量波次", c.get("weak_period_analysis"), [("time", "时间"), ("analysis", "分析"), ("action", "修正动作")])
    document.add_heading("五、话术分别应该承担什么流量任务", level=1)
    for line in ["拉进入后的第一停留：具体身份、单一场景和可回答问题。", "提升停留：可观察细节加一条清晰证明链。", "提升互动与关注：给出低门槛动作，并明确用户能得到什么。", "提升商品点击：证明讲清后再给明确动作，避免只重复库存口令。"]:
        document.add_paragraph(line, style="List Bullet")
    cycle = c.get("cycle_recommendation") if isinstance(c.get("cycle_recommendation"), dict) else {}
    document.add_heading("六、建议的话术循环", level=1)
    for key, label in (("range", "建议时长"), ("reason", "判断依据"), ("allocation", "环节安排")):
        paragraph = document.add_paragraph(); paragraph.add_run(f"{label}：").bold = True; paragraph.add_run(text(cycle.get(key)))
    document.add_heading("七、下一场可以直接执行的改进", level=1)
    for action in items(c.get("host_actions")):
        document.add_paragraph(text(action), style="List Number")
    document.add_heading("八、下场验证指标", level=1)
    for line in ["每轮的进入、离开与净流入", "循环开始、结束和峰值实时在线", "平均观看时长", "评论、点赞和新增关注", "商品曝光、商品点击及点击/曝光", "本轮使用的主场景、个人证明、核心证明和行动表达"]:
        document.add_paragraph(line, style="List Bullet")
    document.core_properties.title = "直播流量与话术关系分析"
    document.core_properties.author = "直播复盘 Agent"
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def generate_outputs(job_dir: Path, session_name: str, sentences: list[dict[str, Any]], timeline: dict[str, Any], review: dict[str, Any], source_inventory: list[dict[str, Any]], settings: Settings) -> list[dict[str, str]]:
    output_dir = job_dir / "output"; output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(c for c in session_name if c not in '<>:"/\\|?*').strip() or "本场直播"
    payload = build_payload(sentences, timeline, review)
    payload_path = job_dir / "workbook_payload.json"
    payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    xlsx_path = output_dir / f"{safe_name}_流量与话术分析.xlsx"
    build_xlsx(xlsx_path, payload)
    docx_path = output_dir / f"{safe_name}_流量与话术分析.docx"
    review_for_docx = dict(review)
    review_for_docx["timeline_minutes"] = timeline.get("minutes", [])
    build_docx(docx_path, review_for_docx)
    note_path = output_dir / "处理说明.txt"
    note_path.write_text("\n".join(["直播流量与话术分析处理说明", "本报告不使用GMV、成交、订单或销量字段。", *review["warnings"]]), encoding="utf-8-sig")
    return [{"label": "下载流量与话术分析 Excel", "path": str(xlsx_path)}, {"label": "下载流量与话术分析 Word", "path": str(docx_path)}, {"label": "下载处理说明", "path": str(note_path)}]
