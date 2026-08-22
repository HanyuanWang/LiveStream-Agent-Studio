from __future__ import annotations

import cgi
import hashlib
import json
import os
import re
import shutil
import socket
import subprocess
import sys
import threading
import time
import uuid
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import quote, unquote, urlparse
from urllib.request import Request, urlopen


ROOT = Path(__file__).resolve().parent
CODEX_ROOT = ROOT.parent
BREAKDOWN = CODEX_ROOT / "live_breakdown_agent"
WORKSPACE = ROOT / "workspace"
JOBS_DIR = WORKSPACE / "jobs"
JOBS: dict[str, dict] = {}
LOCK = threading.Lock()
VALIDATION_FILE = WORKSPACE / "connection_validation.json"
DIRECTOR_DIR = WORKSPACE / "director"
DOUYIN_BROWSER_DIR = DIRECTOR_DIR / "douyin_browser"
DOUYIN_DEBUG_PORT = 9235
ALLOWED_WEB_ORIGINS = {"http://127.0.0.1:4173", "http://localhost:4173"}
ALLOWED_GATEWAY_HOSTS = {"127.0.0.1:8785", "localhost:8785"}


def port_open(port: int) -> bool:
    try:
        with socket.create_connection(("127.0.0.1", port), timeout=.35):
            return True
    except OSError:
        return False


def json_file(path: Path) -> dict:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def safe_name(value: str) -> str:
    return Path(value or "video.mp4").name.replace("\x00", "_")


def parse_env(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    if path.exists():
        for raw in path.read_text(encoding="utf-8-sig", errors="ignore").splitlines():
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def protect_private_config(path: Path) -> None:
    """Best-effort Windows ACL: keep the plaintext compatibility file private."""
    if os.name != "nt" or not path.exists():
        return
    username = os.environ.get("USERNAME", "").strip()
    if not username:
        return
    subprocess.run(
        ["icacls", str(path), "/inheritance:r", "/grant:r", f"{username}:(F)", "SYSTEM:(F)"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        check=False,
        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
    )


def config_fingerprint(values: dict[str, str]) -> str:
    keys = ("DASHSCOPE_API_KEY", "ALIYUN_ACCESS_KEY_ID", "ALIYUN_ACCESS_KEY_SECRET", "ALIYUN_OSS_ENDPOINT", "ALIYUN_OSS_BUCKET")
    return hashlib.sha256("|".join(values.get(k, "") for k in keys).encode()).hexdigest()


def connection_status() -> dict:
    values = parse_env(BREAKDOWN / ".env")
    validation = json_file(VALIDATION_FILE)
    same = validation.get("fingerprint") == config_fingerprint(values)
    model_configured = bool(values.get("DASHSCOPE_API_KEY"))
    oss_configured = all(values.get(k) for k in ("ALIYUN_ACCESS_KEY_ID", "ALIYUN_ACCESS_KEY_SECRET", "ALIYUN_OSS_ENDPOINT", "ALIYUN_OSS_BUCKET"))
    recorder_value = values.get("QUICK_RECORDER_EXE", "").strip()
    recorder_path = Path(recorder_value).expanduser() if recorder_value else None
    return {
        "model": {"provider": "阿里云百炼", "name": values.get("TEXT_MODEL") or "qwen-plus", "configured": model_configured, "verified": bool(same and validation.get("model_verified")), "verified_at": validation.get("verified_at") if same else None},
        "oss": {"provider": "阿里云 OSS", "configured": oss_configured, "verified": bool(same and validation.get("oss_verified")), "endpoint": values.get("ALIYUN_OSS_ENDPOINT", ""), "bucket": values.get("ALIYUN_OSS_BUCKET", ""), "verified_at": validation.get("verified_at") if same else None},
        "recorder": {"configured": bool(recorder_value), "found": bool(recorder_path and recorder_path.is_file()), "path": recorder_value},
    }


def tail_error(path: Path) -> str:
    if not path.exists():
        return "未找到处理日志"
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    high_value = [
        line.strip() for line in lines
        if any(marker in line.lower() for marker in (
            "pip install", "remote end closed", "timed out", "timeout",
            "permission", "access denied", "forbidden", "no such file",
            "not found", "invalid", "quota", "insufficient", "欠费",
        ))
    ]
    if high_value:
        return re.sub(r"^ERROR:\s*", "", high_value[-1]).strip()
    useful = [re.sub(r"^ERROR:\s*", "", x).strip() for x in lines if "ERROR" in x or "错误" in x or "失败" in x]
    specific = [x for x in useful if "错误码" not in x and "处理程序退出" not in x]
    if specific:
        return specific[-1]
    if useful:
        return useful[-1]
    return "；".join(lines[-5:])[-500:] or "处理程序异常退出"


def proxy_json(method: str, path: str, body: bytes | None = None, content_type: str = "application/json") -> tuple[int, bytes, str]:
    request = Request(f"http://127.0.0.1:8765{path}", data=body, method=method, headers={"Content-Type": content_type})
    try:
        with urlopen(request, timeout=120) as response:
            return response.status, response.read(), response.headers.get("Content-Type", "application/json")
    except HTTPError as exc:
        return exc.code, exc.read(), exc.headers.get("Content-Type", "application/json")


def multipart_body(fields: dict[str, str], files: dict[str, Path]) -> tuple[bytes, str]:
    boundary = "----LiveAgent" + uuid.uuid4().hex
    chunks: list[bytes] = []
    for name, value in fields.items():
        chunks += [f"--{boundary}\r\nContent-Disposition: form-data; name=\"{name}\"\r\n\r\n{value}\r\n".encode()]
    for name, path in files.items():
        chunks += [f"--{boundary}\r\nContent-Disposition: form-data; name=\"{name}\"; filename=\"{path.name}\"\r\nContent-Type: application/octet-stream\r\n\r\n".encode(), path.read_bytes(), b"\r\n"]
    chunks += [f"--{boundary}--\r\n".encode()]
    return b"".join(chunks), f"multipart/form-data; boundary={boundary}"


def qwen_json(system: str, user: str) -> dict:
    """Call the configured Qwen model and require a JSON object response."""
    values = parse_env(BREAKDOWN / ".env")
    key = values.get("DASHSCOPE_API_KEY", "").strip()
    if not key:
        raise RuntimeError("尚未配置百炼 API Key，请先到“设置与连接”完成配置和验证")
    model = values.get("TEXT_MODEL") or "qwen-plus"
    payload = json.dumps({
        "model": model,
        "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}],
        "temperature": 0.35,
        "response_format": {"type": "json_object"},
    }, ensure_ascii=False).encode("utf-8")
    request = Request(
        "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
        data=payload,
        method="POST",
        headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
    )
    try:
        with urlopen(request, timeout=120) as response:
            result = json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Qwen 请求失败（HTTP {exc.code}）：{detail[:500]}") from exc
    content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
    if isinstance(content, dict):
        return content
    text = str(content).strip()
    text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.I | re.S).strip()
    try:
        return json.loads(text)
    except Exception as exc:
        match = re.search(r"\{.*\}", text, flags=re.S)
        if match:
            return json.loads(match.group(0))
        raise RuntimeError("Qwen 已返回内容，但不是可解析的 JSON") from exc


def launch_douyin_search(keyword: str) -> None:
    chrome = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")
    if not chrome.exists():
        chrome = Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe")
    if not chrome.exists():
        raise RuntimeError("没有找到 Google Chrome，请先安装 Chrome")
    profile = DIRECTOR_DIR / "douyin_profile"
    profile.mkdir(parents=True, exist_ok=True)
    url = f"https://www.douyin.com/search/{quote(keyword)}?type=video"
    flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    subprocess.Popen(
        [str(chrome), f"--user-data-dir={profile}", "--new-window", url],
        stdin=subprocess.DEVNULL,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        creationflags=flags,
    )


def chrome_path() -> Path:
    for item in (
        Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
        Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
    ):
        if item.exists():
            return item
    raise RuntimeError("没有找到 Google Chrome，请先安装 Chrome")


def _walk_aweme_payload(value):
    if isinstance(value, dict):
        if isinstance(value.get("aweme_info"), dict):
            yield value["aweme_info"]
        if value.get("aweme_id") and (value.get("statistics") or value.get("author") or value.get("desc")):
            yield value
        for child in value.values():
            yield from _walk_aweme_payload(child)
    elif isinstance(value, list):
        for child in value:
            yield from _walk_aweme_payload(child)


def _metric_number(value) -> int:
    try:
        return int(float(value or 0))
    except (TypeError, ValueError):
        return 0


def _aweme_candidate(aweme: dict, keyword: str) -> dict | None:
    aweme_id = str(aweme.get("aweme_id") or aweme.get("id") or "").strip()
    if not aweme_id.isdigit():
        return None
    author = aweme.get("author") if isinstance(aweme.get("author"), dict) else {}
    stats = aweme.get("statistics") if isinstance(aweme.get("statistics"), dict) else {}
    title = str(aweme.get("desc") or aweme.get("preview_title") or "").strip()
    copy_text = title
    video = aweme.get("video") if isinstance(aweme.get("video"), dict) else {}
    captions = video.get("caption_infos") or aweme.get("caption_infos") or []
    caption_parts = []
    if isinstance(captions, list):
        for caption in captions:
            if isinstance(caption, dict):
                text = caption.get("text") or caption.get("content") or caption.get("title")
                if text and str(text).strip() not in caption_parts:
                    caption_parts.append(str(text).strip())
    if caption_parts:
        copy_text = "\n".join(caption_parts)
    published_at = ""
    if aweme.get("create_time"):
        try:
            published_at = datetime.fromtimestamp(int(aweme["create_time"])).strftime("%Y-%m-%d %H:%M:%S")
        except (TypeError, ValueError, OSError):
            pass
    return {
        "aweme_id": aweme_id,
        "url": f"https://www.douyin.com/video/{aweme_id}",
        "title": title or f"抖音视频 {aweme_id}",
        "author": str(author.get("nickname") or author.get("unique_id") or author.get("short_id") or "").strip(),
        "author_id": str(author.get("uid") or author.get("sec_uid") or "").strip(),
        "likes": _metric_number(stats.get("digg_count")),
        "comments": _metric_number(stats.get("comment_count")),
        "shares": _metric_number(stats.get("share_count")),
        "collects": _metric_number(stats.get("collect_count")),
        "published_at": published_at,
        "copy": copy_text or title,
        "search_keyword": keyword,
        "verified": True,
        "source": "抖音搜索页自动读取",
    }


def research_douyin(plan: dict, limit: int = 12) -> dict:
    """Open a dedicated browser, collect real search payloads, and rank by likes."""
    try:
        from playwright.sync_api import sync_playwright  # type: ignore
    except ImportError as exc:
        raise RuntimeError("自动抖音调研组件尚未安装，请重新启动 LiveAgent Studio 完成组件安装") from exc

    keywords = [str(item.get("keyword") or "").strip() for item in (plan.get("keywords") or []) if isinstance(item, dict)]
    keywords = [item for item in keywords if item][:8]
    if not keywords:
        raise RuntimeError("检索方案里没有可用关键词")
    profile = DIRECTOR_DIR / "douyin_auto_profile"
    profile.mkdir(parents=True, exist_ok=True)
    found: dict[str, dict] = {}
    response_count = 0

    with sync_playwright() as playwright:
        context = playwright.chromium.launch_persistent_context(
            str(profile), executable_path=str(chrome_path()), headless=False,
            no_viewport=True, args=["--start-maximized", "--disable-blink-features=AutomationControlled"], timeout=30000,
        )
        page = context.pages[0] if context.pages else context.new_page()
        active_keyword = {"value": keywords[0]}

        def capture(response):
            nonlocal response_count
            if not any(marker in response.url.lower() for marker in ("aweme", "search", "feed")):
                return
            try:
                payload = response.json()
            except Exception:
                return
            response_count += 1
            for aweme in _walk_aweme_payload(payload):
                item = _aweme_candidate(aweme, active_keyword["value"])
                if item and item["copy"]:
                    previous = found.get(item["aweme_id"])
                    if not previous or item["likes"] > previous["likes"]:
                        found[item["aweme_id"]] = item

        page.on("response", capture)
        for keyword in keywords:
            active_keyword["value"] = keyword
            page.goto(f"https://www.douyin.com/search/{quote(keyword)}?type=video", wait_until="domcontentloaded", timeout=45000)
            page.wait_for_timeout(3500)
            body_text = page.locator("body").inner_text(timeout=10000)
            if not found and any(token in body_text for token in ("登录后", "扫码登录", "验证码", "安全验证")):
                deadline = time.time() + 120
                while time.time() < deadline:
                    page.wait_for_timeout(2000)
                    text = page.locator("body").inner_text(timeout=3000)
                    if not any(token in text for token in ("扫码登录", "验证码", "安全验证")):
                        break
            for _ in range(5):
                page.mouse.wheel(0, 1800)
                page.wait_for_timeout(1200)
        context.close()

    ranked = sorted(found.values(), key=lambda item: (item.get("likes", 0), item.get("comments", 0)), reverse=True)
    if not ranked:
        raise RuntimeError("未读取到抖音真实视频数据。请在弹出的专用浏览器完成登录或验证码后，再点击一次自动检索")
    selected = ranked[:max(1, min(limit, 30))]
    return {"candidates": selected, "searched_keywords": keywords, "captured_responses": response_count, "message": f"已自动核验 {len(selected)} 条抖音视频，并按真实点赞数排序"}


def find_ffmpeg() -> Path:
    values = parse_env(BREAKDOWN / ".env")
    candidates = [
        values.get("FFMPEG_PATH", ""),
        str(CODEX_ROOT / "tools" / "ffmpeg" / "bits_unz" / "ffmpeg-master-latest-win64-gpl" / "bin" / "ffmpeg.exe"),
        shutil.which("ffmpeg") or "",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return Path(candidate)
    raise RuntimeError("没有找到 FFmpeg，无法从短视频中提取音轨。请先安装 FFmpeg 或在配置中填写 FFMPEG_PATH")


def format_clock(seconds: float) -> str:
    total = max(0, int(round(float(seconds or 0))))
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}" if hours else f"{minutes:02d}:{secs:02d}"


def extract_source_url(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        raise RuntimeError("这一行没有链接")
    patterns = (
        r"https?://v\.douyin\.com/[A-Za-z0-9_-]+/?",
        r"https?://(?:www\.)?douyin\.com/(?:video|note)/[A-Za-z0-9_-]+(?:\?[^\s\u4e00-\u9fff<>\"']*)?",
        r"https?://www\.iesdouyin\.com/share/(?:video|note)/[A-Za-z0-9_-]+(?:\?[^\s\u4e00-\u9fff<>\"']*)?",
        r"https?://[^\s\u4e00-\u9fff<>\"']+",
    )
    url = ""
    for pattern in patterns:
        matched = re.search(pattern, text, flags=re.IGNORECASE)
        if matched:
            url = matched.group(0).rstrip("，。；、！!？?）)]}>.,;")
            break
    if not url:
        raise RuntimeError("这一行没有识别到完整的 http 或 https 短视频链接")
    return url


def validate_source_url(value: str) -> str:
    url = extract_source_url(value)
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise RuntimeError("请输入完整的 http 或 https 短视频链接")
    hostname = parsed.hostname.lower()
    if hostname in {"localhost", "127.0.0.1", "0.0.0.0", "::1"}:
        raise RuntimeError("不支持本机地址，请填写短视频平台的公开链接")
    if hostname in {"douyin.com", "www.douyin.com"} and parsed.path.rstrip("/") == "":
        raise RuntimeError("分享文案中没有识别到具体视频链接，请重新复制该视频的分享链接")
    return url


def find_chromium_browser() -> Path | None:
    candidates = [
        Path(os.environ.get("PROGRAMFILES", "")) / "Google/Chrome/Application/chrome.exe",
        Path(os.environ.get("PROGRAMFILES(X86)", "")) / "Google/Chrome/Application/chrome.exe",
        Path(os.environ.get("LOCALAPPDATA", "")) / "Google/Chrome/Application/chrome.exe",
        Path(os.environ.get("PROGRAMFILES", "")) / "Microsoft/Edge/Application/msedge.exe",
        Path(os.environ.get("PROGRAMFILES(X86)", "")) / "Microsoft/Edge/Application/msedge.exe",
    ]
    return next((candidate for candidate in candidates if candidate.is_file()), None)


def read_douyin_browser_cookies() -> list[dict]:
    if not port_open(DOUYIN_DEBUG_PORT):
        return []
    try:
        from playwright.sync_api import sync_playwright  # type: ignore
    except ImportError as exc:
        raise RuntimeError("抖音登录组件尚未安装，请重新启动 LiveAgent Studio 完成安装") from exc
    try:
        with sync_playwright() as playwright:
            browser = playwright.chromium.connect_over_cdp(f"http://127.0.0.1:{DOUYIN_DEBUG_PORT}")
            cookies: list[dict] = []
            for context in browser.contexts:
                cookies.extend(context.cookies())
    except Exception as exc:
        raise RuntimeError(f"无法读取 LiveAgent 抖音登录窗口：{str(exc)[:220]}") from exc
    return [
        cookie for cookie in cookies
        if str(cookie.get("domain") or "").lstrip(".").lower().endswith(("douyin.com", "douyinvod.com", "iesdouyin.com"))
        and str(cookie.get("name") or "")
    ]


def douyin_login_status() -> dict:
    running = port_open(DOUYIN_DEBUG_PORT)
    if not running:
        return {"running": False, "ready": False, "logged_in": False, "message": "尚未打开 LiveAgent 抖音登录窗口"}
    try:
        cookies = read_douyin_browser_cookies()
    except Exception as exc:
        return {"running": True, "ready": False, "logged_in": False, "message": str(exc)}
    names = {str(cookie.get("name") or "").lower() for cookie in cookies}
    ready = bool(cookies)
    logged_in = bool(names.intersection({"sessionid", "sessionid_ss", "sid_guard"}))
    message = "抖音登录状态已可用于读取视频" if logged_in else ("已获取新的抖音访问状态；如视频仍受限，请在窗口中登录" if ready else "请在打开的窗口中访问或登录抖音，然后点击刷新状态")
    return {"running": True, "ready": ready, "logged_in": logged_in, "message": message}


def open_douyin_login_browser() -> dict:
    if not port_open(DOUYIN_DEBUG_PORT):
        browser = find_chromium_browser()
        if not browser:
            raise RuntimeError("没有找到 Chrome 或 Edge，无法打开抖音登录窗口")
        DOUYIN_BROWSER_DIR.mkdir(parents=True, exist_ok=True)
        flags = getattr(subprocess, "CREATE_NEW_PROCESS_GROUP", 0) | getattr(subprocess, "DETACHED_PROCESS", 0)
        subprocess.Popen(
            [
                str(browser),
                f"--user-data-dir={DOUYIN_BROWSER_DIR}",
                f"--remote-debugging-port={DOUYIN_DEBUG_PORT}",
                "--remote-debugging-address=127.0.0.1",
                "--no-first-run",
                "--no-default-browser-check",
                "https://www.douyin.com/",
            ],
            cwd=str(browser.parent),
            creationflags=flags,
            close_fds=True,
        )
        for _ in range(30):
            if port_open(DOUYIN_DEBUG_PORT):
                break
            time.sleep(.2)
    status = douyin_login_status()
    status["opened"] = True
    return status


def write_douyin_cookie_file(work_dir: Path) -> Path | None:
    cookies = read_douyin_browser_cookies()
    if not cookies:
        return None
    cookie_path = work_dir / "douyin-cookies.txt"
    lines = ["# Netscape HTTP Cookie File", "# Generated temporarily by LiveAgent Studio"]
    for cookie in cookies:
        domain = str(cookie.get("domain") or "")
        include_subdomains = "TRUE" if domain.startswith(".") else "FALSE"
        path = str(cookie.get("path") or "/")
        secure = "TRUE" if cookie.get("secure") else "FALSE"
        expires = max(0, int(float(cookie.get("expires") or 0)))
        name = str(cookie.get("name") or "").replace("\t", "").replace("\n", "")
        value = str(cookie.get("value") or "").replace("\t", "").replace("\n", "")
        lines.append("\t".join((domain, include_subdomains, path, secure, str(expires), name, value)))
    cookie_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    protect_private_config(cookie_path)
    return cookie_path


def download_director_video(url: str, work_dir: Path) -> tuple[Path, dict]:
    try:
        from yt_dlp import YoutubeDL  # type: ignore
    except ImportError as exc:
        raise RuntimeError("短视频链接解析组件尚未安装，请重新启动 LiveAgent Studio 完成安装") from exc
    work_dir.mkdir(parents=True, exist_ok=True)
    options = {
        "outtmpl": str(work_dir / "source.%(ext)s"),
        "format": "bestaudio/best",
        "noplaylist": True,
        "quiet": True,
        "no_warnings": True,
        "socket_timeout": 45,
        "retries": 2,
        "http_headers": {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/128 Safari/537.36"},
    }
    hostname = (urlparse(url).hostname or "").lower()
    cookie_path: Path | None = None
    if hostname.endswith(("douyin.com", "iesdouyin.com")):
        cookie_path = write_douyin_cookie_file(work_dir)
        if cookie_path:
            options["cookiefile"] = str(cookie_path)
    try:
        with YoutubeDL(options) as downloader:
            info = downloader.extract_info(url, download=True)
            prepared = Path(downloader.prepare_filename(info))
    except Exception as exc:
        reason = str(exc)[:320]
        if "Fresh cookies" in reason or "cookies" in reason.lower():
            raise RuntimeError(f"抖音要求新的访问状态。请点击“登录抖音 / 刷新状态”，在专用窗口中打开或登录抖音后重试。原始错误：{reason}") from exc
        raise RuntimeError(f"平台拒绝读取、链接失效，或该页面需要登录/验证码：{reason}") from exc
    finally:
        if cookie_path:
            try:
                cookie_path.unlink(missing_ok=True)
            except OSError:
                pass
    candidates = [prepared] + sorted(work_dir.glob("source.*"), key=lambda item: item.stat().st_mtime, reverse=True)
    media = next((item for item in candidates if item.exists() and item.suffix.lower() not in {".json", ".part", ".ytdl"}), None)
    if not media:
        raise RuntimeError("链接已读取，但没有找到可用于转写的音视频文件")
    metadata = {
        "title": str(info.get("title") or info.get("fulltitle") or "参考短视频"),
        "author": str(info.get("uploader") or info.get("channel") or info.get("creator") or ""),
        "duration": float(info.get("duration") or 0),
        "url": url,
    }
    return media, metadata


def transcribe_director_source(url: str, work_dir: Path) -> dict:
    status = connection_status()
    if not status["model"]["verified"] or not status["oss"]["verified"]:
        raise RuntimeError("Qwen 或 OSS 尚未完成实际验证，请先到“设置与连接”验证后再处理链接")
    src_dir = BREAKDOWN / "src"
    if str(src_dir) not in sys.path:
        sys.path.insert(0, str(src_dir))
    # 直播拆解 Agent 的 OSS 与语音转写实现都在 aliyun.py。
    # 旧的 asr/oss 模块名已不存在，引用它们会在真正转写前直接失败。
    from live_breakdown_agent.aliyun import DashScopeAsrClient, OssUploader, parse_asr_result  # type: ignore
    from live_breakdown_agent.media import extract_audio  # type: ignore

    values = parse_env(BREAKDOWN / ".env")
    media, metadata = download_director_video(url, work_dir)
    audio_path = work_dir / "audio.flac"
    extract_audio(media, audio_path, find_ffmpeg())
    uploader = OssUploader(
        values["ALIYUN_OSS_ENDPOINT"], values["ALIYUN_OSS_BUCKET"],
        values["ALIYUN_ACCESS_KEY_ID"], values["ALIYUN_ACCESS_KEY_SECRET"],
        values.get("ALIYUN_SECURITY_TOKEN", ""),
    )
    object_key = f"liveagent-director/{uuid.uuid4().hex}/audio.flac"
    try:
        audio_url = uploader.upload_and_sign(audio_path, object_key)
        client = DashScopeAsrClient(
            values["DASHSCOPE_API_KEY"],
            values.get("DASHSCOPE_BASE_URL") or "https://dashscope.aliyuncs.com",
            values.get("ASR_MODEL") or "paraformer-v2",
        )
        task_id = client.submit(audio_url)
        result = client.poll(task_id)
        result_data = client.download_result(result)
        segments = parse_asr_result(result_data)
    finally:
        try:
            uploader.delete(object_key)
        except Exception:
            pass
    if not segments:
        raise RuntimeError("转写服务已经完成，但没有识别出有效语音")
    rows = [{
        "start": float(segment.start), "end": float(segment.end),
        "start_label": format_clock(segment.start), "end_label": format_clock(segment.end),
        "text": str(segment.text).strip(),
    } for segment in segments if str(segment.text).strip()]
    metadata["segments"] = rows
    metadata["transcript"] = "\n".join(f"[{item['start_label']}-{item['end_label']}] {item['text']}" for item in rows)
    metadata["duration_label"] = format_clock(metadata.get("duration") or (rows[-1]["end"] if rows else 0))
    return metadata


def write_director_outputs(output_dir: Path, name: str, brief: dict, sources: list[dict], analysis: dict) -> tuple[Path, Path]:
    from docx import Document  # type: ignore
    from openpyxl import Workbook  # type: ignore
    from openpyxl.styles import Alignment, Font, PatternFill  # type: ignore

    output_dir.mkdir(parents=True, exist_ok=True)
    stem = re.sub(r'[<>:"/\\|?*]+', "_", name).strip() or "视频编导方案"
    word_path = output_dir / f"{stem}_短视频编导方案.docx"
    excel_path = output_dir / f"{stem}_短视频素材与脚本.xlsx"

    doc = Document()
    doc.add_heading(f"{name}｜短视频编导方案", 0)
    doc.add_paragraph("分析范围：只使用用户提供的短视频链接及成功解析出的逐字稿，不判断平台热度，不虚构点赞、评论或成交指标。原创脚本只借鉴内容结构，不复制原视频表达。")
    doc.add_heading("一、创作需求", level=1)
    for key, label in (("product", "产品"), ("category", "赛道"), ("audience", "目标人群"), ("benefits", "核心卖点"), ("price", "价格/活动"), ("pain", "核心痛点"), ("persona", "账号人设"), ("duration", "成片时长"), ("count", "脚本数量"), ("prohibited", "禁用表达")):
        doc.add_paragraph(f"{label}：{brief.get(key) or '未提供'}")
    doc.add_heading("二、参考视频逐字稿", level=1)
    for index, item in enumerate(sources, 1):
        doc.add_heading(f"素材 {index}｜{item.get('title') or item.get('author') or '未命名'}", level=2)
        doc.add_paragraph(f"链接：{item.get('url','')}\n作者：{item.get('author','未识别')}｜时长：{item.get('duration_label','')}")
        doc.add_paragraph(str(item.get("transcript") or ""))
    doc.add_heading("三、内容结构拆解", level=1)
    for item in analysis.get("benchmark_findings", []):
        doc.add_heading(str(item.get("title") or item.get("video") or "素材规律"), level=2)
        doc.add_paragraph(f"有效原因：{item.get('why_it_worked','')}\n开头：{item.get('hook','')}\n结构：{item.get('structure','')}\n可复用模式：{item.get('reusable_pattern','')}\n避免照搬：{item.get('avoid_copy','')}")
    doc.add_heading("四、原创短视频脚本", level=1)
    for index, script in enumerate(analysis.get("scripts", []), 1):
        doc.add_heading(f"脚本 {index}｜{script.get('title','未命名')}", level=2)
        doc.add_paragraph(f"时长：{script.get('duration','')}｜角度：{script.get('angle','')}\n3秒开头：{script.get('hook','')}")
        for shot in script.get("shots", []):
            doc.add_paragraph(f"{shot.get('time','')}｜画面：{shot.get('visual','')}｜口播：{shot.get('voiceover','')}｜字幕：{shot.get('onscreen_text','')}")
        doc.add_paragraph(f"完整口播：{script.get('voiceover','')}\n结尾行动：{script.get('cta','')}\n合规提醒：{script.get('compliance_notes','')}")
    doc.add_heading("五、拍摄执行建议", level=1)
    filming = analysis.get("filming_plan", {})
    for key, value in filming.items():
        doc.add_paragraph(f"{key}：{value if not isinstance(value, list) else '；'.join(map(str,value))}")
    doc.save(word_path)

    wb = Workbook()
    ws = wb.active; ws.title = "结论与脚本"
    headers = ["脚本", "时长", "创意角度", "3秒开头", "完整口播", "CTA", "合规提醒"]
    ws.append(headers)
    for script in analysis.get("scripts", []):
        ws.append([script.get("title"), script.get("duration"), script.get("angle"), script.get("hook"), script.get("voiceover"), script.get("cta"), script.get("compliance_notes")])
    shots = wb.create_sheet("逐镜分镜")
    shots.append(["脚本", "镜头序号", "时间", "画面", "口播", "字幕"])
    for script in analysis.get("scripts", []):
        for index, shot in enumerate(script.get("shots", []), 1):
            shots.append([script.get("title"), index, shot.get("time"), shot.get("visual"), shot.get("voiceover"), shot.get("onscreen_text")])
    material = wb.create_sheet("参考视频逐字稿")
    material.append(["素材", "作者", "链接", "开始", "结束", "逐字稿"])
    for item in sources:
        for segment in item.get("segments", []):
            material.append([item.get("title"), item.get("author"), item.get("url"), segment.get("start_label"), segment.get("end_label"), segment.get("text")])
    findings = wb.create_sheet("内容拆解")
    findings.append(["主题", "有效原因", "开头", "结构", "可复用模式", "避免照搬"])
    for item in analysis.get("benchmark_findings", []):
        findings.append([item.get("title"), item.get("why_it_worked"), item.get("hook"), item.get("structure"), item.get("reusable_pattern"), item.get("avoid_copy")])
    filming_sheet = wb.create_sheet("拍摄建议")
    filming_sheet.append(["项目", "建议"])
    for key, value in filming.items():
        filming_sheet.append([key, "；".join(map(str, value)) if isinstance(value, list) else value])
    for sheet in wb.worksheets:
        sheet.freeze_panes = "A2"
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="0D6C43")
        for column in sheet.columns:
            letter = column[0].column_letter
            sheet.column_dimensions[letter].width = min(50, max(12, max(len(str(c.value or "")) for c in column) + 2))
            for cell in column:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.save(excel_path)
    return word_path, excel_path


def save_job(job_id: str) -> None:
    target = JOBS_DIR / job_id / "state.json"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(JOBS[job_id], ensure_ascii=False, indent=2), encoding="utf-8")


def update_job(job_id: str, **changes) -> None:
    with LOCK:
        JOBS[job_id].update(changes)
        save_job(job_id)


def monitor_breakdown(
    job_id: str,
    process: subprocess.Popen,
    started: datetime,
    before: dict[Path, float] | None = None,
    requested_output_dir: Path | None = None,
) -> None:
    status_path = BREAKDOWN / "workspace" / "current_status.json"
    output_dir = BREAKDOWN / "workspace" / "output"
    before = before if before is not None else ({p.resolve(): p.stat().st_mtime for p in output_dir.glob("*.xlsx")} if output_dir.exists() else {})
    while process.poll() is None:
        state = json_file(status_path)
        message = state.get("message") or "正在提取音频、上传并转写"
        stage = state.get("state", "running")
        progress_map = {"starting": 8, "running": 38, "audio_uploaded": 45, "transcription_submitted": 55, "completed": 100}
        current = JOBS.get(job_id, {}).get("progress", 5)
        progress = max(current, progress_map.get(stage, min(92, current + 1)))
        update_job(job_id, status="running", status_text="处理中", progress=progress, message=message)
        threading.Event().wait(1.2)
    if process.returncode == 0:
        outputs = sorted(
            (p for p in output_dir.glob("*.xlsx") if p.resolve() not in before or p.stat().st_mtime > before[p.resolve()]),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
        if not outputs:
            update_job(job_id, status="failed", status_text="处理失败", progress=100, message="处理程序已结束，但没有找到新生成的拆解 Excel", error="没有找到新生成的拆解 Excel", finished_at=datetime.now().isoformat(timespec="seconds"))
            return
        generated = outputs[0]
        result = generated
        try:
            if requested_output_dir:
                requested_output_dir.mkdir(parents=True, exist_ok=True)
                destination = requested_output_dir / generated.name
                if destination.resolve() != generated.resolve():
                    shutil.copy2(generated, destination)
                result = destination
        except Exception as exc:
            update_job(job_id, status="failed", status_text="保存失败", progress=100, message=f"拆解已完成，但复制到指定目录失败：{exc}", error=f"复制结果失败：{exc}", source_output=str(generated), finished_at=datetime.now().isoformat(timespec="seconds"))
            return
        update_job(job_id, status="completed", status_text="已完成", progress=100, message="直播拆解已完成", output=str(result), source_output=str(generated), finished_at=datetime.now().isoformat(timespec="seconds"))
    else:
        own_log = Path(JOBS[job_id].get("log") or "")
        shared_log = BREAKDOWN / "workspace" / "logs" / "latest.log"
        reason = tail_error(own_log if own_log.exists() else shared_log)
        update_job(job_id, status="failed", status_text="处理失败", progress=100, message=reason, error=reason, error_log=str(own_log if own_log.exists() else shared_log), finished_at=datetime.now().isoformat(timespec="seconds"))


def monitor_review(job_id: str, review_job_id: str, requested_output_dir: Path | None = None) -> None:
    while True:
        try:
            with urlopen(f"http://127.0.0.1:8775/api/jobs/{review_job_id}", timeout=5) as response:
                state = json.loads(response.read().decode("utf-8"))
            status = state.get("status", "running")
            outputs = state.get("outputs") or []
            if status == "completed":
                source_dir = CODEX_ROOT / "live_retro_agent" / "workspace" / "jobs" / review_job_id / "output"
                if requested_output_dir:
                    requested_output_dir.mkdir(parents=True, exist_ok=True)
                local_outputs = []
                for entry in outputs:
                    name = str(entry.get("name") or Path(unquote(str(entry.get("url") or ""))).name)
                    source = source_dir / name
                    if not source.exists():
                        continue
                    target = source
                    if requested_output_dir:
                        target = requested_output_dir / source.name
                        if target.resolve() != source.resolve():
                            shutil.copy2(source, target)
                    local_outputs.append({**entry, "name": source.name, "path": str(target), "url": str(target)})
                workbooks = [entry for entry in local_outputs if str(entry.get("path", "")).lower().endswith(".xlsx")]
                primary = (workbooks or local_outputs)[0].get("path") if (workbooks or local_outputs) else None
                update_job(job_id, status=status, status_text=state.get("status_text", "已完成"), progress=100, message=state.get("message") or "Excel、Word 和处理说明均已生成", output=primary, output_dir=str(requested_output_dir or source_dir), review_job_id=review_job_id, outputs=local_outputs)
            else:
                update_job(job_id, status=status, status_text=state.get("status_text", "分析中"), progress=int(state.get("progress") or 5), message=state.get("message") or "正在分析直播话术与流量变化", review_job_id=review_job_id, outputs=outputs)
            if status in {"completed", "failed"}:
                return
        except Exception as exc:
            update_job(job_id, status="failed", status_text="处理失败", progress=100, message=f"复盘服务连接失败：{exc}", error=str(exc))
            return
        threading.Event().wait(1.5)


class Handler(BaseHTTPRequestHandler):
    server_version = "LiveAgentStudioGateway/0.2"

    def log_message(self, fmt, *args):
        return

    def cors(self) -> None:
        origin = self.headers.get("Origin", "")
        if origin in ALLOWED_WEB_ORIGINS:
            self.send_header("Access-Control-Allow-Origin", origin)
            self.send_header("Vary", "Origin")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, DELETE, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def trusted_local_request(self) -> bool:
        host = self.headers.get("Host", "").lower()
        origin = self.headers.get("Origin", "")
        return host in ALLOWED_GATEWAY_HOSTS and (not origin or origin in ALLOWED_WEB_ORIGINS)

    def reject_untrusted_request(self) -> bool:
        if self.trusted_local_request():
            return False
        self.send_json({"error": "请求来源不受信任；本机接口只接受 LiveAgent Studio 页面调用"}, 403)
        return True

    def send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False, default=str).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Cache-Control", "no-store")
        self.cors()
        self.end_headers()
        self.wfile.write(body)

    def do_OPTIONS(self):
        if self.reject_untrusted_request():
            return
        self.send_response(204)
        self.cors()
        self.end_headers()

    def do_GET(self):
        if self.reject_untrusted_request():
            return
        path = urlparse(self.path).path
        if path == "/api/health":
            connections = connection_status()
            self.send_json({
                "ok": True,
                "mode": "local",
                "services": {"scout": port_open(8765), "review": port_open(8775), "gateway": True},
                **connections,
            })
            return
        if path == "/api/director/douyin-login":
            self.send_json(douyin_login_status())
            return
        if path == "/api/jobs":
            items=[]
            for state_path in JOBS_DIR.glob("*/state.json"):
                state=json_file(state_path)
                if not state: continue
                if state.get("status")=="failed":
                    log=Path(state.get("log") or state.get("error_log") or "")
                    reason=tail_error(log)
                    state.update({"message":reason,"error":reason})
                    state_path.write_text(json.dumps(state,ensure_ascii=False,indent=2),encoding="utf-8")
                items.append(state)
            items.sort(key=lambda x:x.get("created_at", ""),reverse=True)
            self.send_json({"jobs":items})
            return
        if path.startswith("/api/jobs/"):
            job_id = path.rsplit("/", 1)[-1]
            state = JOBS.get(job_id)
            if not state:
                state = json_file(JOBS_DIR / job_id / "state.json")
            self.send_json(state or {"error": "任务不存在"}, 200 if state else 404)
            return
        if path.startswith("/api/scout/"):
            proxied = path.removeprefix("/api/scout") or "/api/status"
            try:
                status, body, content_type = proxy_json("GET", proxied)
                self.send_response(status); self.send_header("Content-Type", content_type); self.send_header("Content-Length", str(len(body))); self.cors(); self.end_headers(); self.wfile.write(body)
            except Exception as exc:
                self.send_json({"error": f"主播发现服务不可用：{exc}"}, 503)
            return
        self.send_json({"error": "Not found"}, 404)

    def do_POST(self):
        if self.reject_untrusted_request():
            return
        path = urlparse(self.path).path
        if path == "/api/director/douyin-login":
            try:
                self.send_json(open_douyin_login_browser())
            except Exception as exc:
                self.send_json({"error": str(exc)}, 500)
            return
        if path == "/api/director/sources":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                data = json.loads(self.rfile.read(length).decode("utf-8") if length else "{}")
                raw_links = data.get("links") or []
                links, failures = [], []
                for value in raw_links:
                    text = str(value or "").strip()
                    if not text:
                        continue
                    try:
                        url = validate_source_url(text)
                        if url not in links:
                            links.append(url)
                    except Exception as exc:
                        failures.append({"url": text[:240], "error": str(exc)[:600]})
                links = links[:10]
                if not links:
                    self.send_json({"error": "没有识别到可处理的短视频链接", "failures": failures}, 400); return
                status = connection_status()
                if not status["model"]["verified"] or not status["oss"]["verified"]:
                    self.send_json({"error": "Qwen 或 OSS 尚未完成实际验证，请先到“设置与连接”完成验证"}, 400); return
                sources = []
                for url in links:
                    work_dir = DIRECTOR_DIR / "sources" / uuid.uuid4().hex
                    try:
                        sources.append(transcribe_director_source(url, work_dir))
                    except Exception as exc:
                        failures.append({"url": url, "error": str(exc)[:600]})
                    finally:
                        shutil.rmtree(work_dir, ignore_errors=True)
                self.send_json({
                    "sources": sources,
                    "failures": failures,
                    "message": f"成功生成 {len(sources)} 条逐字稿；{len(failures)} 条失败",
                })
            except Exception as exc:
                self.send_json({"error": f"短视频链接处理失败：{exc}"}, 500)
            return
        if path == "/api/director/jobs":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                data = json.loads(self.rfile.read(length).decode("utf-8") if length else "{}")
                brief = data.get("brief") or {}
                sources = data.get("sources") or []
                verified = [item for item in sources if str(item.get("url") or "").strip() and str(item.get("transcript") or "").strip()]
                if not verified:
                    self.send_json({"error": "请先解析至少一个短视频链接并生成逐字稿"}, 400); return
                requested_text = str(data.get("output_dir") or "").strip().strip('"')
                requested_output_dir = Path(requested_text).expanduser() if requested_text else None
                if requested_output_dir and not requested_output_dir.is_absolute():
                    self.send_json({"error": "保存位置必须是完整路径"}, 400); return
                if requested_output_dir: requested_output_dir.mkdir(parents=True, exist_ok=True)
                job_id = datetime.now().strftime("%Y%m%d_%H%M%S_") + uuid.uuid4().hex[:7]
                state = {"job_id": job_id, "name": str(data.get("name") or f"{brief.get('product','产品')}视频编导方案"), "agent": "视频编导", "status": "running", "status_text": "分析素材", "progress": 8, "message": "已收到参考逐字稿，正在拆解内容结构", "created_at": datetime.now().isoformat(timespec="seconds"), "output_dir": str(requested_output_dir or (DIRECTOR_DIR / "output" / job_id))}
                with LOCK: JOBS[job_id] = state; save_job(job_id)
                threading.Thread(target=self._director_pipeline, args=(job_id, brief, verified, requested_output_dir), daemon=True).start()
                self.send_json(state, 201)
            except Exception as exc:
                self.send_json({"error": str(exc)}, 500)
            return
        if path == "/api/select-folder":
            try:
                script = (
                    "Add-Type -AssemblyName System.Windows.Forms; "
                    "Add-Type -TypeDefinition 'using System; using System.Runtime.InteropServices; "
                    "public static class LiveAgentForeground { "
                    "[DllImport(\"user32.dll\")] public static extern bool SetForegroundWindow(IntPtr hWnd); }'; "
                    "$owner=New-Object System.Windows.Forms.Form; "
                    "$owner.Text='LiveAgent Studio'; "
                    "$owner.TopMost=$true; $owner.ShowInTaskbar=$false; "
                    "$owner.StartPosition=[System.Windows.Forms.FormStartPosition]::CenterScreen; "
                    "$owner.Width=1; $owner.Height=1; $owner.Opacity=0; "
                    "$owner.Show(); $owner.Activate(); "
                    "[LiveAgentForeground]::SetForegroundWindow($owner.Handle)|Out-Null; "
                    "$d=New-Object System.Windows.Forms.FolderBrowserDialog; "
                    "$d.Description='选择任务输出文件的保存文件夹'; "
                    "$d.ShowNewFolderButton=$true; "
                    "$result=$d.ShowDialog($owner); $owner.Close(); $owner.Dispose(); "
                    "if($result -eq [System.Windows.Forms.DialogResult]::OK){"
                    "[Console]::OutputEncoding=[System.Text.UTF8Encoding]::new(); $d.SelectedPath}; "
                    "$d.Dispose()"
                )
                creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
                completed = subprocess.run(
                    ["powershell.exe", "-NoProfile", "-STA", "-Command", script],
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    timeout=300,
                    creationflags=creationflags,
                )
                selected = completed.stdout.strip().splitlines()[-1].strip() if completed.stdout.strip() else ""
                self.send_json({"path": selected, "cancelled": not bool(selected)})
            except Exception as exc:
                self.send_json({"error": f"无法打开文件夹选择器：{exc}"}, 500)
            return
        if path == "/api/connections/configure":
            try:
                length=int(self.headers.get("Content-Length","0")); data=json.loads(self.rfile.read(length).decode("utf-8") if length else "{}")
                env_path=BREAKDOWN/".env"; values=parse_env(env_path); previous_fingerprint=config_fingerprint(values)
                mapping={"dashscope_api_key":"DASHSCOPE_API_KEY","oss_access_key_id":"ALIYUN_ACCESS_KEY_ID","oss_access_key_secret":"ALIYUN_ACCESS_KEY_SECRET","oss_endpoint":"ALIYUN_OSS_ENDPOINT","oss_bucket":"ALIYUN_OSS_BUCKET"}
                for field,key in mapping.items():
                    value=str(data.get(field) or "").strip()
                    if "\n" in value or "\r" in value:
                        raise ValueError(f"{field} 不能包含换行符")
                    if value: values[key]=value
                recorder_path = str(data.get("quick_recorder_exe") or "").strip().strip('"')
                if recorder_path:
                    candidate = Path(recorder_path).expanduser()
                    if not candidate.is_absolute():
                        raise ValueError("录制助手路径必须是完整的绝对路径")
                    if candidate.suffix.lower() != ".exe":
                        raise ValueError("录制助手路径必须指向 .exe 文件")
                    if not candidate.is_file():
                        raise ValueError("没有找到该录制助手 EXE，请检查粘贴的路径")
                    values["QUICK_RECORDER_EXE"] = str(candidate)
                else:
                    values.pop("QUICK_RECORDER_EXE", None)
                lines=[f"{key}={value}" for key,value in values.items()]; env_path.write_text("\n".join(lines)+"\n",encoding="utf-8"); protect_private_config(env_path)
                if config_fingerprint(values) != previous_fingerprint and VALIDATION_FILE.exists(): VALIDATION_FILE.unlink()
                self.send_json(connection_status())
            except Exception as exc: self.send_json({"error":str(exc)},400)
            return
        if path == "/api/connections/verify":
            values = parse_env(BREAKDOWN / ".env")
            model_ok = False; oss_ok = False; errors = []
            try:
                key = values.get("DASHSCOPE_API_KEY", "")
                request = Request("https://dashscope.aliyuncs.com/compatible-mode/v1/models", headers={"Authorization": f"Bearer {key}"})
                with urlopen(request, timeout=12) as response: model_ok = response.status == 200
            except Exception as exc: errors.append(f"百炼验证失败：{exc}")
            try:
                import oss2  # type: ignore
                auth = oss2.Auth(values.get("ALIYUN_ACCESS_KEY_ID", ""), values.get("ALIYUN_ACCESS_KEY_SECRET", ""))
                bucket = oss2.Bucket(auth, values.get("ALIYUN_OSS_ENDPOINT", ""), values.get("ALIYUN_OSS_BUCKET", ""))
                bucket.get_bucket_info(); oss_ok = True
            except Exception as exc: errors.append(f"OSS验证失败：{exc}")
            result = {"fingerprint": config_fingerprint(values), "model_verified": model_ok, "oss_verified": oss_ok, "verified_at": datetime.now().isoformat(timespec="seconds"), "errors": errors}
            VALIDATION_FILE.parent.mkdir(parents=True, exist_ok=True); VALIDATION_FILE.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
            self.send_json({**connection_status(), "errors": errors}, 200 if model_ok and oss_ok else 400)
            return
        if path.startswith("/api/scout/"):
            proxied = path.removeprefix("/api/scout")
            length = int(self.headers.get("Content-Length", "0")); body = self.rfile.read(length) if length else b"{}"
            try:
                status, response_body, content_type = proxy_json("POST", proxied, body, self.headers.get("Content-Type", "application/json"))
                self.send_response(status); self.send_header("Content-Type", content_type); self.send_header("Content-Length", str(len(response_body))); self.cors(); self.end_headers(); self.wfile.write(response_body)
            except Exception as exc: self.send_json({"error": f"主播发现服务不可用：{exc}"}, 503)
            return
        if path == "/api/breakdown/jobs":
            try:
                form = cgi.FieldStorage(fp=self.rfile, headers=self.headers, environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": self.headers.get("Content-Type", ""), "CONTENT_LENGTH": self.headers.get("Content-Length", "0")})
                item = form["video"] if "video" in form else None
                if item is None or not getattr(item, "filename", ""):
                    self.send_json({"error": "请选择直播视频"}, 400)
                    return
                job_id = datetime.now().strftime("%Y%m%d_%H%M%S_") + uuid.uuid4().hex[:7]
                requested_text = str(getattr(form["output_dir"], "value", "") if "output_dir" in form else "").strip().strip('"')
                requested_output_dir = Path(requested_text).expanduser() if requested_text else None
                if requested_output_dir and not requested_output_dir.is_absolute():
                    self.send_json({"error": "保存位置必须是完整路径，例如 C:\\Users\\你的名字\\Desktop\\拆解结果"}, 400)
                    return
                if requested_output_dir:
                    try:
                        requested_output_dir.mkdir(parents=True, exist_ok=True)
                    except Exception as exc:
                        self.send_json({"error": f"无法使用该保存位置：{exc}"}, 400)
                        return
                input_dir = JOBS_DIR / job_id / "input"
                input_dir.mkdir(parents=True, exist_ok=True)
                target = input_dir / safe_name(item.filename)
                with target.open("wb") as handle:
                    shutil.copyfileobj(item.file, handle)
                python = Path(sys.executable)
                python_cmd = str(python) if python.exists() else "python"
                env = os.environ.copy()
                env["LIVE_AGENT_NO_POPUP"] = "1"
                env["LIVE_AGENT_NO_PAUSE"] = "1"
                env["PYTHONIOENCODING"] = "utf-8"
                env["PYTHONPATH"] = str(BREAKDOWN / "src") + (os.pathsep + env["PYTHONPATH"] if env.get("PYTHONPATH") else "")
                log_path = JOBS_DIR / job_id / "process.log"
                log = log_path.open("w", encoding="utf-8")
                native_output_dir = BREAKDOWN / "workspace" / "output"
                before = {p.resolve(): p.stat().st_mtime for p in native_output_dir.glob("*.xlsx")} if native_output_dir.exists() else {}
                process = subprocess.Popen([python_cmd, str(BREAKDOWN / "process_video_with_status.py"), str(target)], cwd=BREAKDOWN, env=env, stdout=log, stderr=subprocess.STDOUT)
                task_name = str(getattr(form["name"], "value", "") if "name" in form else "").strip() or target.stem
                state = {"job_id": job_id, "name": task_name, "agent": "直播拆解", "status": "running", "status_text": "正在启动", "progress": 3, "message": "视频已收到，正在检查配置", "created_at": datetime.now().isoformat(timespec="seconds"), "video": str(target), "log": str(log_path), "output_dir": str(requested_output_dir or native_output_dir)}
                with LOCK:
                    JOBS[job_id] = state
                    save_job(job_id)
                threading.Thread(target=monitor_breakdown, args=(job_id, process, datetime.now(), before, requested_output_dir), daemon=True).start()
                self.send_json(state, 201)
            except Exception as exc:
                self.send_json({"error": str(exc)}, 500)
            return
        if path == "/api/review/jobs":
            try:
                form = cgi.FieldStorage(fp=self.rfile, headers=self.headers, environ={"REQUEST_METHOD":"POST","CONTENT_TYPE":self.headers.get("Content-Type", ""),"CONTENT_LENGTH":self.headers.get("Content-Length", "0")})
                video = form["video"] if "video" in form else None; traffic = form["traffic"] if "traffic" in form else None
                if video is None or not getattr(video,"filename","") or traffic is None or not getattr(traffic,"filename",""):
                    self.send_json({"error":"请同时上传直播视频和巨量百应流量数据"},400); return
                requested_text = str(getattr(form["output_dir"], "value", "") if "output_dir" in form else "").strip().strip('"')
                requested_output_dir = Path(requested_text).expanduser() if requested_text else None
                if requested_output_dir and not requested_output_dir.is_absolute():
                    self.send_json({"error":"保存位置必须是完整路径，例如 C:\\Users\\你的名字\\Desktop\\复盘结果"},400); return
                if requested_output_dir:
                    try: requested_output_dir.mkdir(parents=True,exist_ok=True)
                    except Exception as exc: self.send_json({"error":f"无法使用该保存位置：{exc}"},400); return
                job_id = datetime.now().strftime("%Y%m%d_%H%M%S_")+uuid.uuid4().hex[:7]; input_dir=JOBS_DIR/job_id/"input"; input_dir.mkdir(parents=True,exist_ok=True)
                video_path=input_dir/safe_name(video.filename); traffic_path=input_dir/safe_name(traffic.filename)
                with video_path.open("wb") as handle: shutil.copyfileobj(video.file,handle)
                with traffic_path.open("wb") as handle: shutil.copyfileobj(traffic.file,handle)
                state={"job_id":job_id,"name":str(getattr(form["name"],"value","") or video_path.stem),"agent":"直播复盘","status":"running","status_text":"生成逐字稿","progress":3,"message":"视频已收到，正在先生成逐字稿","created_at":datetime.now().isoformat(timespec="seconds"),"video":str(video_path),"traffic":str(traffic_path),"output_dir":str(requested_output_dir or "")}
                with LOCK: JOBS[job_id]=state; save_job(job_id)
                threading.Thread(target=self._review_pipeline,args=(job_id,video_path,traffic_path,requested_output_dir),daemon=True).start(); self.send_json(state,201)
            except Exception as exc: self.send_json({"error":str(exc)},500)
            return
        match = re.fullmatch(r"/api/jobs/([^/]+)/(open-file|open-folder)", path)
        if match:
            job_id, action = match.groups()
            state = JOBS.get(job_id) or json_file(JOBS_DIR / job_id / "state.json")
            output_value = str(state.get("output") or "") if state else ""
            if not output_value:
                self.send_json({"error": "该任务还没有生成输出文件"}, 404)
                return
            review_url = re.fullmatch(r"/api/jobs/([^/]+)/files/(.+)", output_value)
            if review_url:
                output = CODEX_ROOT / "live_retro_agent" / "workspace" / "jobs" / review_url.group(1) / "output" / unquote(review_url.group(2))
            else:
                output = Path(output_value)
            if not output.exists():
                self.send_json({"error": f"输出文件不存在或已被移动：{output}"}, 404)
                return
            # Older tasks stored only the shared output directory. Resolve the
            # most likely workbook so "open file" also works for those jobs.
            if output.is_dir() and action == "open-file":
                task_name = str(state.get("name") or "").strip().lower()
                workbooks = sorted(
                    output.glob("*.xlsx"),
                    key=lambda candidate: candidate.stat().st_mtime,
                    reverse=True,
                )
                matched = [
                    candidate
                    for candidate in workbooks
                    if task_name and task_name in candidate.stem.lower()
                ]
                if matched:
                    output = matched[0]
                elif workbooks:
                    output = workbooks[0]
                else:
                    self.send_json({"error": "输出目录中没有找到 Excel 结果文件"}, 404)
                    return
            target = output if action == "open-file" else (output if output.is_dir() else output.parent)
            try:
                os.startfile(str(target))
                self.send_json({"ok": True, "path": str(target)})
            except Exception as exc:
                self.send_json({"error": f"无法打开：{exc}"}, 500)
            return
        if path.startswith("/api/open-output/"):
            agent = path.rsplit("/", 1)[-1]
            targets = {"breakdown": BREAKDOWN / "workspace" / "output", "review": CODEX_ROOT / "live_retro_agent" / "workspace" / "output", "scout": CODEX_ROOT / "live_scout_agent" / "workspace", "director": DIRECTOR_DIR / "output"}
            target = targets.get(agent)
            if target:
                target.mkdir(parents=True, exist_ok=True)
                os.startfile(target)
                self.send_json({"ok": True})
            else:
                self.send_json({"error": "未知 Agent"}, 404)
            return
        self.send_json({"error": "Not found"}, 404)

    @staticmethod
    def _director_pipeline(job_id: str, brief: dict, sources: list[dict], requested_output_dir: Path | None = None) -> None:
        try:
            update_job(job_id, progress=28, status_text="拆解文案", message="正在比较开头、痛点、卖点、信任和行动引导")
            payload = {"brief": brief, "reference_sources": sources}
            analysis = qwen_json(
                "你是资深短视频编导。只使用用户提供链接中成功转写出的真实逐字稿作为参考，不搜索其他视频，不判断高赞，不虚构页面指标。提炼结构和表达策略，但不得照抄原视频措辞；随后按用户需求生成可拍摄的原创脚本、逐镜分镜和拍摄建议。涉及保健品、食品、美妆等领域时主动规避医疗功效、绝对化和无法证实的承诺。只输出JSON。",
                "输入：" + json.dumps(payload, ensure_ascii=False) + "\n输出：{\"executive_summary\":\"\",\"benchmark_findings\":[{\"title\":\"\",\"why_it_worked\":\"\",\"hook\":\"\",\"structure\":\"\",\"reusable_pattern\":\"\",\"avoid_copy\":\"\"}],\"scripts\":[{\"title\":\"\",\"duration\":\"\",\"angle\":\"\",\"hook\":\"\",\"shots\":[{\"time\":\"\",\"visual\":\"\",\"voiceover\":\"\",\"onscreen_text\":\"\"}],\"voiceover\":\"\",\"cta\":\"\",\"compliance_notes\":\"\"}],\"filming_plan\":{\"location\":[],\"people\":[],\"props\":[],\"camera\":[],\"performance\":[],\"editing\":[]},\"production_checklist\":[],\"editing_handoff\":[]}。脚本数量服从brief.count，每条必须能直接拍摄。",
            )
            update_job(job_id, progress=78, status_text="生成方案", message="文案规律已完成，正在生成原创脚本和拍摄建议")
            output_dir = requested_output_dir or (DIRECTOR_DIR / "output" / job_id)
            word_path, excel_path = write_director_outputs(output_dir, JOBS[job_id]["name"], brief, sources, analysis)
            outputs = [{"name": word_path.name, "path": str(word_path), "url": str(word_path)}, {"name": excel_path.name, "path": str(excel_path), "url": str(excel_path)}]
            update_job(job_id, status="completed", status_text="已完成", progress=100, message="短视频编导方案、原创脚本和拍摄建议已生成", output=str(word_path), outputs=outputs, output_dir=str(output_dir), finished_at=datetime.now().isoformat(timespec="seconds"))
        except Exception as exc:
            update_job(job_id, status="failed", status_text="处理失败", progress=100, message=f"视频编导生成失败：{exc}", error=str(exc), finished_at=datetime.now().isoformat(timespec="seconds"))

    @staticmethod
    def _review_pipeline(job_id: str, video_path: Path, traffic_path: Path, requested_output_dir: Path | None = None) -> None:
        python = Path(sys.executable); env=os.environ.copy(); env.update({"LIVE_AGENT_NO_POPUP":"1","LIVE_AGENT_NO_PAUSE":"1","PYTHONIOENCODING":"utf-8"}); env["PYTHONPATH"] = str(BREAKDOWN / "src") + (os.pathsep + env["PYTHONPATH"] if env.get("PYTHONPATH") else "")
        log_path=JOBS_DIR/job_id/"breakdown.log"
        with log_path.open("w",encoding="utf-8") as log: process=subprocess.Popen([str(python),str(BREAKDOWN/"process_video_with_status.py"),str(video_path)],cwd=BREAKDOWN,env=env,stdout=log,stderr=subprocess.STDOUT)
        while process.poll() is None: update_job(job_id,progress=min(58,JOBS[job_id].get("progress",3)+1),message="正在从视频生成逐句逐字稿"); threading.Event().wait(1.5)
        if process.returncode!=0: reason=tail_error(log_path); update_job(job_id,status="failed",status_text="处理失败",progress=100,message=reason,error=reason,error_log=str(log_path)); return
        output_dir=BREAKDOWN/"workspace"/"output"; candidates=sorted(output_dir.glob(f"{video_path.stem}*_拆解.xlsx"),key=lambda p:p.stat().st_mtime,reverse=True)
        if not candidates: update_job(job_id,status="failed",status_text="处理失败",progress=100,message="逐字稿已经处理，但没有找到拆解 Excel"); return
        update_job(job_id,progress=62,status_text="大模型分析",message="逐字稿已完成，正在与巨量百应流量时间线对齐")
        body,content_type=multipart_body({"session_name":JOBS[job_id]["name"]},{"breakdown":candidates[0],"minute":traffic_path})
        try:
            request=Request("http://127.0.0.1:8775/api/jobs",data=body,method="POST",headers={"Content-Type":content_type})
            with urlopen(request,timeout=120) as response: review_job_id=json.loads(response.read().decode("utf-8"))["job_id"]
            monitor_review(job_id,review_job_id,requested_output_dir)
        except Exception as exc: update_job(job_id,status="failed",status_text="处理失败",progress=100,message=f"提交大模型分析失败：{exc}",error=str(exc))

    def do_DELETE(self):
        if self.reject_untrusted_request():
            return
        path=urlparse(self.path).path
        if path.startswith("/api/jobs/"):
            job_id=path.rsplit("/",1)[-1]
            with LOCK: JOBS.pop(job_id,None)
            target=JOBS_DIR/job_id
            if target.exists(): shutil.rmtree(target)
            self.send_json({"ok":True}); return
        if path.startswith("/api/scout/"):
            proxied=path.removeprefix("/api/scout")
            try:
                status,response_body,content_type=proxy_json("DELETE",proxied,b"",self.headers.get("Content-Type","application/json"))
                self.send_response(status); self.send_header("Content-Type",content_type); self.send_header("Content-Length",str(len(response_body))); self.cors(); self.end_headers(); self.wfile.write(response_body)
            except Exception as exc:self.send_json({"error":f"主播发现服务不可用：{exc}"},503)
            return
        self.send_json({"error":"Not found"},404)


def main() -> None:
    JOBS_DIR.mkdir(parents=True, exist_ok=True)
    for stale_cookie in DIRECTOR_DIR.glob("**/douyin-cookies.txt"):
        try:
            stale_cookie.unlink(missing_ok=True)
        except OSError:
            pass
    server = ThreadingHTTPServer(("127.0.0.1", 8785), Handler)
    print("LiveAgent Studio local gateway: http://127.0.0.1:8785", flush=True)
    server.serve_forever()


if __name__ == "__main__":
    main()
