import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from live_scout_agent.creator_reports import (
    CreatorReportManager,
    build_report_docx,
    fallback_analysis,
    finalize_commerce_analysis,
    validate_detail_evidence,
)


class CreatorReportTests(unittest.TestCase):
    def test_builds_editable_report_with_core_sections(self):
        candidate = {
            "id": 1,
            "anchor_name": "测试达人",
            "douyin_id": "test001",
            "category": "服饰内衣",
            "followers": 86_000,
            "estimated_gmv": 750_000,
            "sales_volume": 15_000,
            "sessions_7d": 3,
            "source": "蝉妈妈",
            "reasons": ["低粉高转化"],
            "raw_data": {"销售客单价": "200~300"},
        }
        analysis = finalize_commerce_analysis(
            candidate,
            {
                **fallback_analysis(candidate),
                "period_summary": {
                    "period": "近30天",
                    "total_gmv": 750_000,
                    "session_count": 3,
                    "sales_aov": "200~300",
                },
                "sessions": [
                    {
                        "date": "2026-07-01 10:00",
                        "gmv": 250_000,
                        "sales_volume": 5_000,
                        "aov": 50,
                        "duration": "3小时",
                        "products": "测试商品A",
                    },
                    {
                        "date": "2026-07-08 10:00",
                        "gmv": 200_000,
                        "sales_volume": 4_000,
                        "aov": 50,
                        "duration": "3小时",
                        "products": "测试商品B",
                    },
                    {
                        "date": "2026-07-15 10:00",
                        "gmv": 300_000,
                        "sales_volume": 6_000,
                        "aov": 50,
                        "duration": "3小时",
                        "products": "测试商品A、测试商品B",
                    },
                ],
                "product_breakdown": [
                    {
                        "product_name": "测试商品A",
                        "category": "连衣裙",
                        "product_url": "https://example.com/product-a",
                        "gmv": 450_000,
                        "sales_volume": 9_000,
                        "session_count": 2,
                    },
                    {
                        "product_name": "测试商品B",
                        "category": "针织衫",
                        "gmv": 300_000,
                        "sales_volume": 6_000,
                        "session_count": 2,
                    },
                ],
            },
        )
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "report.docx"
            build_report_docx(candidate, analysis, path)
            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            )
        self.assertIn("测试达人深度分析报告", text)
        self.assertIn("人设定位", text)
        self.assertIn("带货与商品策略", text)
        self.assertIn("近30天/当前可见周期场次表现", text)
        self.assertIn("商品结构拆解", text)
        self.assertIn("场次GMV", table_text)
        self.assertIn("GMV占比", table_text)
        self.assertIn("测试商品A", table_text)
        self.assertIn("同行可复用的关键打法", text)
        self.assertGreaterEqual(len(document.tables), 4)
        self.assertEqual(len(document.sections), 3)
        self.assertEqual(document.sections[1].orientation, WD_ORIENT.LANDSCAPE)
        for table in document.tables[1:3]:
            header_properties = table.rows[0]._tr.get_or_add_trPr()
            self.assertIsNotNone(header_properties.find(qn("w:tblHeader")))

    def test_computes_product_gmv_share_and_session_stability(self):
        candidate = {"estimated_gmv": 1_000_000, "raw_data": {"销售客单价": "100~200"}}
        result = finalize_commerce_analysis(
            candidate,
            {
                "sessions": [
                    {"date": "2026-07-01", "gmv": 400_000},
                    {"date": "2026-07-02", "gmv": 600_000},
                ],
                "product_breakdown": [
                    {"product_name": "商品A", "gmv": 250_000, "sales_volume": 1_000},
                    {"product_name": "商品B", "gmv": 750_000, "sales_volume": 1_500},
                ],
            },
        )
        self.assertEqual(result["period_summary"]["sales_aov"], "100~200")
        self.assertAlmostEqual(result["product_breakdown"][0]["gmv_share"], 25.0)
        self.assertAlmostEqual(result["product_breakdown"][0]["aov"], 250.0)
        self.assertAlmostEqual(result["period_summary"]["stability"]["average_gmv"], 500_000)
        self.assertAlmostEqual(result["period_summary"]["stability"]["max_share"], 60.0)

    def test_llm_indexes_cannot_be_used_as_actual_sales_or_gmv(self):
        candidate = {
            "estimated_gmv": None,
            "estimated_gmv_text": "1000w+",
            "gmv_index": 5_609_284,
            "sales_volume": None,
            "sales_volume_text": "25w~50w",
            "sales_index": 51_273,
        }
        result = finalize_commerce_analysis(
            candidate,
            {
                "_numeric_source": "llm",
                "period_summary": {
                    "period": "近30天",
                    "total_gmv": 6_510_501,
                    "total_gmv_range": "1000w+",
                },
                "sessions": [
                    {
                        "date": "2026-07-28",
                        "gmv": 831_000,
                        "gmv_is_actual": True,
                        "gmv_range": "50w~100w",
                        "gmv_index": 831_000,
                        "sales_volume": 7_309,
                        "sales_volume_is_actual": True,
                        "sales_volume_range": "2.5w~5w",
                        "sales_index": 7_309,
                    }
                ],
                "metrics": [
                    ["近30天总GMV", "6510501"],
                    ["近30天总销量", "61297"],
                    ["销售客单价", "100~200"],
                ],
            },
        )
        session = result["sessions"][0]
        self.assertIsNone(result["period_summary"]["total_gmv"])
        self.assertEqual(result["period_summary"]["total_gmv_text"], "1000w+")
        self.assertIsNone(session["gmv"])
        self.assertEqual(session["gmv_range"], "50w~100w")
        self.assertIsNone(session["sales_volume"])
        self.assertEqual(session["sales_volume_range"], "2.5w~5w")
        self.assertIsNone(result["period_summary"]["stability"]["average_gmv"])
        metric_labels = [row[0] for row in result["metrics"]]
        self.assertNotIn("近30天总GMV", metric_labels)
        self.assertNotIn("近30天总销量", metric_labels)
        self.assertIn("销售客单价", metric_labels)
        self.assertTrue(any("销量指数未参与" in note for note in result["data_notes"]))

    def test_flags_session_count_mismatch_without_fabricating_rows(self):
        result = finalize_commerce_analysis(
            {"sessions_7d": 11},
            {
                "period_summary": {"period": "近30天", "session_count": 11},
                "sessions": [{"date": f"2026-07-{day:02d}", "gmv": 100_000} for day in range(1, 11)],
            },
        )
        self.assertEqual(len(result["sessions"]), 10)
        self.assertEqual(len(result["period_summary"]["data_conflicts"]), 1)
        self.assertIn("汇总显示直播场次为11", result["period_summary"]["data_conflicts"][0])

    def test_report_contains_complete_leaderboard_without_web_appendix(self):
        candidate = {
            "anchor_name": "完整数据达人",
            "douyin_id": "all-data-001",
            "estimated_gmv_text": "1000w+",
            "gmv_index": 5_609_284,
            "sales_volume_text": "25w~50w",
            "sales_index": 51_273,
            "followers": 262_000,
            "category": "服饰内衣",
            "sessions_7d": 10,
            "raw_data": {"排行": 11, "销售客单价": "100~200", "直播间": "直播"},
        }
        analysis = fallback_analysis(candidate)
        analysis["web_snapshot"] = {
            "source": "蝉妈妈达人详情页",
            "captured_at": "2026-07-31T12:00:00+08:00",
            "analysis_url": "https://example.com/creator",
            "leaderboard": {"rank": 11},
            "pages": [
                {
                    "section": "基础分析",
                    "tab": "数据概览",
                    "title": "达人详情",
                    "url": "https://example.com/creator/basic",
                    "captured_at": "2026-07-31T12:00:01+08:00",
                    "period": "30天",
                    "lines": ["粉丝总数 26.2w", "活跃粉丝数/占比 4.4w / 17.0%"],
                    "tables": [
                        {
                            "headers": ["日期", "销量", "销售额"],
                            "rows": [["2026-07-30", "2.5w~5w", "500w~1000w"]],
                        }
                    ],
                    "chart_labels": ["销售额趋势"],
                }
            ],
            "warnings": [],
        }
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "all-data-report.docx"
            build_report_docx(candidate, analysis, path)
            document = Document(path)
            text = "\n".join(
                [paragraph.text for paragraph in document.paragraphs]
                + [
                    cell.text
                    for table in document.tables
                    for row in table.rows
                    for cell in row.cells
                ]
            )
        self.assertIn("榜单原始数据（完整字段）", text)
        self.assertIn("销售额指数（仅供比较）", text)
        self.assertIn("销量指数（仅供比较）", text)
        self.assertNotIn("蝉妈妈账号详情完整数据附录", text)
        self.assertNotIn("基础分析 / 数据概览", text)
        self.assertNotIn("粉丝总数 26.2w", text)

    def test_rejects_report_when_detail_page_was_not_collected(self):
        with self.assertRaisesRegex(RuntimeError, "没有采集到"):
            validate_detail_evidence(
                {
                    "analysis_url": "https://example.com/creator",
                    "pages": [],
                }
            )
        with self.assertRaisesRegex(RuntimeError, "未成功进入"):
            validate_detail_evidence(
                {
                    "analysis_url": "https://example.com/creator",
                    "pages": [
                        {
                            "section": "账号概览",
                            "lines": ["榜单数据"],
                        }
                    ],
                }
            )

    def test_manager_overwrites_canonical_report_after_complete_collection(self):
        evidence = {
            "analysis_url": "https://example.com/creator",
            "pages": [
                {
                    "section": "基础分析",
                    "tab": "数据概览",
                    "lines": ["粉丝数 10万"],
                    "tables": [],
                    "chart_labels": [],
                }
            ],
        }

        class FakeDatabase:
            def update_candidate_status(self, candidate_ids, status):
                return len(candidate_ids)

        with tempfile.TemporaryDirectory() as directory:
            report_dir = Path(directory)
            canonical = report_dir / "测试达人_达人拆解报告.docx"
            canonical.write_bytes(b"old-report")
            settings = SimpleNamespace(report_dir=report_dir)
            manager = CreatorReportManager(settings, FakeDatabase())
            candidate = {
                "id": 1,
                "anchor_name": "测试达人",
                "analysis_url": "https://example.com/creator",
            }

            def fake_build(candidate, analysis, output_path):
                output_path.write_bytes(b"new-report")
                return output_path

            with (
                patch(
                    "live_scout_agent.creator_reports.collect_chanmama_evidence",
                    return_value=evidence,
                ),
                patch(
                    "live_scout_agent.creator_reports.analyze_creator",
                    return_value={"web_snapshot": evidence},
                ),
                patch(
                    "live_scout_agent.creator_reports.build_report_docx",
                    side_effect=fake_build,
                ),
            ):
                manager._run([candidate])

            self.assertEqual(canonical.read_bytes(), b"new-report")
            self.assertFalse(list(report_dir.glob("*_20*.docx")))
            evidence_path = report_dir / "测试达人_蝉妈妈网页完整数据.json"
            self.assertTrue(evidence_path.exists())
            self.assertIn("基础分析", evidence_path.read_text(encoding="utf-8"))


if __name__ == "__main__":
    unittest.main()
