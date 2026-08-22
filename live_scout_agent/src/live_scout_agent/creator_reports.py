from __future__ import annotations

import json
import re
import statistics
import threading
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote, urljoin

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Cm, Pt, RGBColor

from .config import Settings
from .chanmama import CHANMAMA_CDP_PORT
from .database import Database
from .http_client import request_json
from .importers import parse_number


ACCENT = "FF6B4A"
INK = "172033"
MUTED = "7D8799"
LIGHT = "F5F7FA"
GREEN = "19B979"
AMBER = "F4B740"


def safe_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", value).strip(" .")
    return cleaned or "达人"


def format_metric(value: Any) -> str:
    if value in (None, ""):
        return "数据不足"
    if isinstance(value, (int, float)):
        number = float(value)
        if abs(number) >= 100_000_000:
            return f"{number / 100_000_000:.2f}亿"
        if abs(number) >= 10_000:
            return f"{number / 10_000:.1f}万"
        return f"{number:,.0f}"
    return str(value)


def _clean_json(text: str) -> dict[str, Any]:
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip())
    return json.loads(cleaned)


def _evidence_payload(evidence: Any) -> dict[str, Any]:
    if isinstance(evidence, dict):
        return evidence
    if isinstance(evidence, str) and evidence.strip():
        try:
            parsed = json.loads(evidence)
            return parsed if isinstance(parsed, dict) else {"legacy_text": evidence}
        except json.JSONDecodeError:
            return {"legacy_text": evidence}
    return {}


def _compact_evidence_for_ai(evidence: Any, limit: int = 120_000) -> str:
    payload = _evidence_payload(evidence)
    if not payload:
        return str(evidence or "")[:limit]
    compact = {
        "source": payload.get("source"),
        "captured_at": payload.get("captured_at"),
        "leaderboard": payload.get("leaderboard"),
        "pages": [],
        "warnings": payload.get("warnings") or [],
    }
    used = len(json.dumps(compact, ensure_ascii=False, default=str))
    for page in payload.get("pages") or []:
        if not isinstance(page, dict):
            continue
        record = {
            "section": page.get("section"),
            "tab": page.get("tab"),
            "title": page.get("title"),
            "url": page.get("url"),
            "period": page.get("period"),
            "lines": (page.get("lines") or [])[:500],
            "tables": page.get("tables") or [],
            "access_note": page.get("access_note"),
        }
        encoded = json.dumps(record, ensure_ascii=False, default=str)
        if used + len(encoded) > limit:
            break
        compact["pages"].append(record)
        used += len(encoded)
    return json.dumps(compact, ensure_ascii=False, default=str)


DETAIL_SECTIONS = (
    "基础分析",
    "带货分析",
    "直播分析",
    "视频分析",
    "粉丝分析",
    "投放分析",
    "推荐诊断",
)

SECTION_TABS = {
    "基础分析": (
        "数据概览", "品类", "一级品类", "二级品类", "三级品类", "品牌", "小店",
        "带货表现", "账号动态", "销售数据", "相似达人",
    ),
    "带货分析": (
        "数据概览", "带货趋势", "30天带货最佳", "带货分布",
        "品类", "一级品类", "二级品类", "三级品类", "品牌", "小店",
        "价格带分布", "商品记录",
    ),
    "直播分析": ("数据概览", "直播带货趋势", "直播流量结构", "直播转化趋势", "直播表现分析", "直播记录"),
    "视频分析": ("数据概览", "视频指标趋势", "视频发布行为", "视频记录"),
    "粉丝分析": (
        "粉丝趋势", "粉丝价值分析", "近30天活跃粉丝趋势", "粉丝团趋势",
        "粉丝画像", "账号粉丝", "视频观众", "直播观众", "画像概览",
        "省", "市", "一级分类", "二级分类",
    ),
    "投放分析": ("直播推广", "商品推广", "直播投放", "商品投放"),
    "推荐诊断": ("相似达人", "综合相似", "视频相似", "直播带货诊断", "视频带货诊断"),
}


def validate_detail_evidence(evidence: Any) -> dict[str, Any]:
    payload = _evidence_payload(evidence)
    pages = [page for page in payload.get("pages") or [] if isinstance(page, dict)]
    detail_pages = [
        page
        for page in pages
        if str(page.get("section") or "").strip() in DETAIL_SECTIONS
    ]
    if not payload.get("analysis_url"):
        raise RuntimeError("未取得蝉妈妈达人详情页地址，不能生成达人拆解报告")
    if not pages:
        raise RuntimeError("没有采集到蝉妈妈达人详情页数据，不能生成残缺报告")
    if not detail_pages:
        raise RuntimeError(
            "未成功进入达人账号的数据分析界面；请重新登录蝉妈妈后再生成报告"
        )
    nonempty_pages = [
        page
        for page in detail_pages
        if page.get("lines") or page.get("tables") or page.get("chart_labels")
    ]
    if not nonempty_pages:
        raise RuntimeError("达人详情页各栏目均未取得可见数据，不能生成残缺报告")
    return payload


def _raw_candidate_data(candidate: dict[str, Any]) -> dict[str, Any]:
    value = candidate.get("raw_data")
    if isinstance(value, dict):
        return value
    if isinstance(value, str) and value.strip():
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except json.JSONDecodeError:
            return {}
    return {}


def _candidate_aov(candidate: dict[str, Any]) -> str:
    for key in ("sales_aov", "sales_aov_text", "aov"):
        if candidate.get(key) not in (None, ""):
            return str(candidate[key])
    raw = _raw_candidate_data(candidate)
    for key in ("销售客单价", "客单价", "场均客单价"):
        if raw.get(key) not in (None, ""):
            return str(raw[key])
    return "数据不足"


def _number(value: Any) -> float | None:
    parsed = parse_number(value)
    if parsed is None or parsed < 0:
        return None
    return float(parsed)


def _range_text(value: Any) -> str:
    if not isinstance(value, str):
        return ""
    text = value.strip()
    return text if re.search(r"[-~～至+]|\b以上\b", text) else ""


def _candidate_metric_text(candidate: dict[str, Any], exact_key: str, text_key: str) -> str:
    text = str(candidate.get(text_key) or "").strip()
    if text:
        return text
    return format_metric(candidate.get(exact_key))


def _normalize_display_units(value: str) -> str:
    text = re.sub(r"(?i)(\d+(?:\.\d+)?)\s*w\b", r"\1万", value)
    return re.sub(r"(?i)(\d+(?:\.\d+)?)\s*k\b", r"\1千", text)


def _format_quantity(exact: Any, range_value: Any = "") -> str:
    range_text = _range_text(range_value)
    return _normalize_display_units(range_text) if range_text else format_metric(exact)


def _format_index(value: Any) -> str:
    number = _number(value)
    return f"{number:,.0f}" if number is not None else "数据不足"


def _format_money_metric(exact: Any, range_value: Any = "") -> str:
    range_text = _range_text(range_value)
    if range_text:
        return "¥" + _normalize_display_units(range_text.replace("￥", "").replace("¥", ""))
    return _format_currency(exact)


def _format_currency(value: Any) -> str:
    number = _number(value)
    if number is None:
        return "数据不足"
    return "¥" + format_metric(number)


def _format_price(value: Any) -> str:
    if isinstance(value, str) and re.search(r"[-~～至]", value):
        text = value.strip().replace("￥", "").replace("¥", "")
        return "¥" + text
    return _format_currency(value)


def _format_percent(value: Any) -> str:
    number = _number(value)
    if number is None:
        return "数据不足"
    if 0 < number <= 1:
        number *= 100
    return f"{number:.1f}%"


def _normalize_session_rows(rows: Any, *, trust_exact: bool = True) -> list[dict[str, Any]]:
    if not isinstance(rows, list):
        return []
    normalized: list[dict[str, Any]] = []
    for item in rows:
        if not isinstance(item, dict):
            continue
        gmv_range = _range_text(item.get("gmv_range") or item.get("gmv_text") or item.get("gmv"))
        sales_range = _range_text(
            item.get("sales_volume_range") or item.get("sales_range") or item.get("sales_volume")
        )
        gmv = _number(item.get("gmv")) if trust_exact else None
        sales_volume = (
            _number(item.get("sales_volume") or item.get("sales"))
            if trust_exact
            else None
        )
        normalized.append(
            {
                "date": str(item.get("date") or item.get("start_time") or "时间未取得"),
                "title": str(item.get("title") or ""),
                "gmv": gmv,
                "gmv_range": gmv_range,
                "gmv_index": _number(item.get("gmv_index")),
                "sales_volume": sales_volume,
                "sales_volume_range": sales_range,
                "sales_index": _number(item.get("sales_index")),
                "order_count": _number(item.get("order_count") or item.get("orders")),
                "aov": item.get("aov"),
                "duration": str(item.get("duration") or "数据不足"),
                "products": str(item.get("products") or item.get("main_products") or "数据不足"),
                "source": str(item.get("source") or "蝉妈妈可见详情页"),
            }
        )
    return normalized


def _normalize_product_rows(
    rows: Any,
    period_gmv: float | None,
    *,
    trust_exact: bool = True,
) -> list[dict[str, Any]]:
    if not isinstance(rows, list):
        return []
    normalized: list[dict[str, Any]] = []
    for item in rows:
        if not isinstance(item, dict):
            continue
        gmv_range = _range_text(item.get("gmv_range") or item.get("gmv_text") or item.get("gmv"))
        sales_range = _range_text(
            item.get("sales_volume_range") or item.get("sales_range") or item.get("sales_volume")
        )
        gmv = _number(item.get("gmv")) if trust_exact else None
        sales_volume = (
            _number(item.get("sales_volume") or item.get("sales"))
            if trust_exact
            else None
        )
        aov = _number(item.get("aov") or item.get("average_price"))
        if aov is None and gmv is not None and sales_volume:
            aov = gmv / sales_volume
        share = (
            _number(item.get("gmv_share") or item.get("share"))
            if trust_exact or item.get("gmv_share_is_actual") is True
            else None
        )
        if share is not None and 0 < share <= 1:
            share *= 100
        if share is None and gmv is not None and period_gmv:
            share = gmv / period_gmv * 100
        normalized.append(
            {
                "product_name": str(item.get("product_name") or item.get("name") or "商品名未取得"),
                "category": str(item.get("category") or item.get("subcategory") or "数据不足"),
                "product_url": str(item.get("product_url") or item.get("url") or ""),
                "aov": aov,
                "gmv": gmv,
                "gmv_range": gmv_range,
                "gmv_index": _number(item.get("gmv_index")),
                "gmv_share": share,
                "sales_volume": sales_volume,
                "sales_volume_range": sales_range,
                "sales_index": _number(item.get("sales_index")),
                "session_count": _number(item.get("session_count") or item.get("sessions")),
                "source": str(item.get("source") or "蝉妈妈可见详情页"),
            }
        )
    return normalized


def _session_stability(rows: list[dict[str, Any]]) -> dict[str, Any]:
    values = [float(row["gmv"]) for row in rows if row.get("gmv") is not None]
    if not values:
        return {
            "session_count": len(rows),
            "gmv_session_count": 0,
            "total_gmv": None,
            "average_gmv": None,
            "median_gmv": None,
            "max_gmv": None,
            "max_share": None,
            "coefficient_of_variation": None,
            "label": "缺少逐场GMV，无法判断稳定性",
        }
    total = sum(values)
    average = statistics.mean(values)
    coefficient = statistics.pstdev(values) / average if len(values) > 1 and average else 0.0
    max_share = max(values) / total * 100 if total else 0.0
    if len(values) < 2:
        label = "仅取得1场GMV，暂不能判断稳定性"
    elif max_share >= 60 or coefficient >= 1:
        label = "单场依赖明显，成交稳定性较弱"
    elif max_share >= 40 or coefficient >= 0.6:
        label = "场次表现波动较大，需要继续观察"
    else:
        label = "逐场GMV相对均衡，成交稳定性较好"
    return {
        "session_count": len(rows),
        "gmv_session_count": len(values),
        "total_gmv": total,
        "average_gmv": average,
        "median_gmv": statistics.median(values),
        "max_gmv": max(values),
        "max_share": max_share,
        "coefficient_of_variation": coefficient,
        "label": label,
    }


def finalize_commerce_analysis(candidate: dict[str, Any], analysis: dict[str, Any]) -> dict[str, Any]:
    result = dict(analysis)
    trust_exact = result.pop("_numeric_source", "") != "llm"
    period_summary = result.get("period_summary")
    if not isinstance(period_summary, dict):
        period_summary = {}
    period_summary = dict(period_summary)
    period_summary.setdefault("period", "近30天/当前可见榜单周期")
    period_gmv_text = (
        _range_text(period_summary.get("total_gmv_range"))
        or _range_text(period_summary.get("total_gmv"))
        or str(candidate.get("estimated_gmv_text") or "").strip()
    )
    period_gmv = _number(period_summary.get("total_gmv")) if trust_exact else None
    if period_gmv is None:
        period_gmv = _number(candidate.get("estimated_gmv"))
    period_summary["total_gmv"] = period_gmv
    period_summary["total_gmv_text"] = period_gmv_text
    period_summary.setdefault("session_count", _number(candidate.get("sessions_7d")))
    period_summary.setdefault("sales_aov", _candidate_aov(candidate))
    period_summary.setdefault("source", "蝉妈妈榜单与可见详情页")

    sessions = _normalize_session_rows(result.get("sessions"), trust_exact=trust_exact)
    products = _normalize_product_rows(
        result.get("product_breakdown"),
        period_gmv,
        trust_exact=trust_exact,
    )
    stability = _session_stability(sessions)
    period_summary["stability"] = stability

    notes = [str(item) for item in (result.get("data_notes") or []) if str(item).strip()]
    conflicts = [
        str(item)
        for item in (period_summary.get("data_conflicts") or [])
        if str(item).strip()
    ]
    if sessions:
        notes = [
            item
            for item in notes
            if not item.startswith(("未取得逐场直播明细", "已识别"))
        ]
    if not sessions:
        notes.append("未取得逐场直播明细，无法验证每场GMV及稳定性")
    elif stability["gmv_session_count"] < len(sessions):
        notes.append(
            f"已识别{len(sessions)}场直播，但只有{stability['gmv_session_count']}场取得GMV"
        )
    declared_session_count = _number(period_summary.get("session_count"))
    if sessions and declared_session_count is not None and int(declared_session_count) != len(sessions):
        conflict = (
            f"页面汇总显示直播场次为{int(declared_session_count)}，"
            f"当前取得逐场记录{len(sessions)}条；两者不一致，报告不据此补造缺失场次"
        )
        conflicts.append(conflict)
        notes.append(conflict)
    if any(row.get("sales_volume_range") for row in sessions):
        notes.append("场次销量按蝉妈妈页面原始区间展示；销量指数未参与销量、订单或客单价计算")
    if any(row.get("gmv_range") for row in sessions):
        notes.append("场次GMV按蝉妈妈页面原始区间展示；销售额指数未参与GMV统计")
    if products:
        notes = [
            item
            for item in notes
            if not item.startswith(("未取得商品级GMV明细", "部分商品缺少GMV"))
        ]
    if not products:
        notes.append("未取得商品级GMV明细，无法计算各商品GMV占比")
    elif any(item.get("gmv_share") is None for item in products):
        notes.append("部分商品缺少GMV或统计周期总GMV，对应GMV占比无法计算")
    period_summary["data_conflicts"] = list(dict.fromkeys(conflicts))

    metrics = [
        list(row)
        for row in (result.get("metrics") or [])
        if isinstance(row, (list, tuple)) and len(row) >= 2
    ]
    if not trust_exact:
        unsafe_metric_markers = ("GMV", "销售额", "成交额", "销量", "销售量", "订单")
        metrics = [
            row
            for row in metrics
            if not any(marker in str(row[0]) for marker in unsafe_metric_markers)
            or bool(_range_text(row[1]))
        ]
    metric_indexes = {str(row[0]): index for index, row in enumerate(metrics)}
    additions = [
        ("统计周期", period_summary["period"]),
        ("周期GMV", _format_money_metric(period_gmv, period_gmv_text)),
        ("直播场次", format_metric(period_summary.get("session_count"))),
        ("销售客单价", str(period_summary.get("sales_aov") or "数据不足")),
        ("场均GMV", _format_currency(stability.get("average_gmv"))),
        ("GMV中位数", _format_currency(stability.get("median_gmv"))),
    ]
    for label, value in additions:
        if label in metric_indexes:
            metrics[metric_indexes[label]][1] = value
        else:
            metrics.append([label, value])

    # Normalize common headline metrics that the model may return as raw
    # database floats (for example ``262000.0`` or ``11.0``).  The underlying
    # values stay unchanged in the separately saved evidence JSON; only the
    # report display is made human-readable.
    display_overrides = {
        "粉丝数": format_metric(candidate.get("followers")),
        "近30天直播场次": format_metric(period_summary.get("session_count")),
    }
    for row in metrics:
        label = str(row[0])
        if label in display_overrides:
            row[1] = display_overrides[label]

    commerce_strategy = [
        str(item)
        for item in (result.get("commerce_strategy") or [])
        if str(item).strip()
    ]
    if products:
        commerce_strategy = [
            item
            for item in commerce_strategy
            if not (item.startswith("主要商品：") and "数据不足" in item)
        ]
        top_products = sorted(
            products,
            key=lambda item: float(item.get("gmv") or 0),
            reverse=True,
        )
        commerce_strategy.insert(
            0,
            "已验证商品："
            + "、".join(str(item["product_name"]) for item in top_products[:8]),
        )
    risks = [str(item) for item in (result.get("risks") or []) if str(item).strip()]
    if products:
        risks = [item for item in risks if not item.startswith("核心商品缺少")]

    result.update(
        period_summary=period_summary,
        sessions=sessions,
        product_breakdown=products,
        metrics=metrics,
        commerce_strategy=commerce_strategy,
        risks=risks,
        data_notes=list(dict.fromkeys(notes)),
    )
    return result


def fallback_analysis(candidate: dict[str, Any], evidence: Any = "") -> dict[str, Any]:
    metrics = [
        ["粉丝数", format_metric(candidate.get("followers"))],
        ["直播销售额", _format_money_metric(candidate.get("estimated_gmv"), candidate.get("estimated_gmv_text"))],
        ["直播销量", _format_quantity(candidate.get("sales_volume"), candidate.get("sales_volume_text"))],
        ["直播场次", format_metric(candidate.get("sessions_7d"))],
        ["销售客单价", _candidate_aov(candidate)],
        ["GPM", format_metric(candidate.get("gpm"))],
        ["UV价值", format_metric(candidate.get("uv_value"))],
    ]
    strengths = list(candidate.get("reasons") or [])
    if candidate.get("estimated_gmv"):
        strengths.append("榜单期内具备可观的直播成交表现")
    gaps = []
    for label, key in (
        ("内容策略", "title"),
        ("核心商品", "products"),
        ("粉丝画像", "audience"),
    ):
        if not candidate.get(key):
            gaps.append(f"{label}缺少可验证数据")
    if not evidence:
        gaps.append("未取得蝉妈妈达人详情页数据，本报告仅依据榜单字段生成")
    analysis = {
        "subtitle": "基于蝉妈妈榜单与可见详情数据的结构化分析",
        "positioning": [
            f"账号：{candidate.get('anchor_name') or '未知'}",
            f"抖音号：{candidate.get('douyin_id') or '数据不足'}",
            f"带货类目：{candidate.get('category') or '数据不足'}",
            f"账号类型：{candidate.get('account_type') or '数据不足'}",
        ],
        "metrics": metrics,
        "persona": ["现有数据不足以可靠判断完整人设，需结合主页简介与内容样本复核"],
        "content_strategy": [
            f"直播/内容标题：{candidate.get('title') or '数据不足'}",
            "需补充近30天视频主题、发布频率、爆款样本和互动数据后判断内容打法",
        ],
        "commerce_strategy": [
            f"主要商品：{candidate.get('products') or '数据不足'}",
            f"榜单直播销售额：{_format_money_metric(candidate.get('estimated_gmv'), candidate.get('estimated_gmv_text'))}",
            f"榜单直播销量：{_format_quantity(candidate.get('sales_volume'), candidate.get('sales_volume_text'))}",
        ],
        "audience": ["当前导出榜单未提供完整性别、年龄、地区和城市等级分布"],
        "recent_changes": ["需要至少两个连续周期数据才能识别增长、回落或异常峰值"],
        "strengths": strengths or ["暂无足够证据形成确定性亮点"],
        "risks": gaps or ["暂未发现明确风险，仍需补充内容与粉丝数据"],
        "cooperation": [
            "先以小规模试播或单品测试验证真实转化",
            "合作前核对类目匹配、历史客单价、退货率与自然流量占比",
        ],
        "reusable_playbook": [
            "按周持续跟踪销售额、销量、粉丝增量和直播场次",
            "把高成交场次与对应商品、内容主题和流量来源进行交叉复盘",
        ],
        "period_summary": {
            "period": "当前可见榜单周期",
            "total_gmv": candidate.get("estimated_gmv"),
            "total_gmv_range": candidate.get("estimated_gmv_text"),
            "session_count": candidate.get("sessions_7d"),
            "sales_aov": _candidate_aov(candidate),
            "source": candidate.get("source") or "蝉妈妈榜单",
        },
        "sessions": [],
        "product_breakdown": [],
        "data_notes": gaps or ["结论仅使用已提供及蝉妈妈页面可见数据"],
    }
    finalized = finalize_commerce_analysis(candidate, analysis)
    finalized["web_snapshot"] = _evidence_payload(evidence)
    return finalized


def analyze_creator(
    candidate: dict[str, Any],
    evidence: Any,
    settings: Settings,
) -> dict[str, Any]:
    fallback = fallback_analysis(candidate, evidence)
    if not settings.dashscope_api_key:
        return fallback
    schema = {
        "subtitle": "一句话说明分析范围",
        "positioning": ["账号定位结论"],
        "metrics": [["指标", "数值"]],
        "persona": ["人设与账号属性"],
        "content_strategy": ["内容主题、形式、频率、爆款与互动"],
        "commerce_strategy": ["带货模式、品类、价格、商品与直播节奏"],
        "audience": ["性别、年龄、地区、城市等级、活跃度"],
        "recent_changes": ["近期增长、回落、异常峰值及可能原因"],
        "strengths": ["有数据依据的亮点"],
        "risks": ["短板或风险"],
        "cooperation": ["品牌、分销、供应链、投放等合作潜力"],
        "reusable_playbook": ["同行可以复用的具体打法"],
        "period_summary": {
            "period": "页面明确显示的统计周期，例如近30天",
            "total_gmv": None,
            "total_gmv_range": "页面原始GMV区间；没有则为null",
            "gmv_index": None,
            "session_count": None,
            "sales_aov": "页面原始客单价或价格区间",
            "source": "数值所在页面或榜单",
        },
        "sessions": [
            {
                "date": "开播日期或时间",
                "title": "场次标题",
                "gmv": None,
                "gmv_range": "页面原始场次GMV区间；没有则为null",
                "gmv_index": None,
                "gmv_is_actual": False,
                "sales_volume": None,
                "sales_volume_range": "页面原始销量区间；没有则为null",
                "sales_index": None,
                "sales_volume_is_actual": False,
                "order_count": None,
                "aov": None,
                "duration": "直播时长",
                "products": "本场明确出现的主要商品",
                "source": "数值所在页面",
            }
        ],
        "product_breakdown": [
            {
                "product_name": "商品名称",
                "category": "细分品类",
                "product_url": "页面明确提供的商品链接",
                "aov": None,
                "gmv": None,
                "gmv_range": "页面原始商品GMV区间；没有则为null",
                "gmv_index": None,
                "gmv_is_actual": False,
                "gmv_share": None,
                "gmv_share_is_actual": False,
                "sales_volume": None,
                "sales_volume_range": "页面原始销量区间；没有则为null",
                "sales_index": None,
                "sales_volume_is_actual": False,
                "session_count": None,
                "source": "数值所在页面",
            }
        ],
        "data_notes": ["缺失数据和结论边界"],
    }
    prompt = {
        "candidate": candidate,
        "visible_chanmama_evidence": _compact_evidence_for_ai(evidence),
    }
    try:
        response = request_json(
            "POST",
            f"{settings.dashscope_base_url}/compatible-mode/v1/chat/completions",
            {"Authorization": f"Bearer {settings.dashscope_api_key}"},
            {
                "model": settings.text_model,
                "temperature": 0.1,
                "response_format": {"type": "json_object"},
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "你是直播电商达人研究Agent。只能使用输入证据，不得编造。"
                            "输出严格JSON，结构与示例完全一致。缺失信息必须写入data_notes。"
                            "重点提取近30天或页面明确显示周期内的每场直播GMV，以及商品名称、"
                            "细分品类、商品链接、商品GMV、销量和客单价。金额统一输出人民币元的"
                            "纯数字；百分比输出0到100的数字。页面没有的数据必须用null，绝不能"
                            "把30天总GMV平均分配到每场，也不能根据话术猜测商品GMV。"
                            "必须严格区分页面大字显示的GMV/销量区间与灰色小字指数："
                            "区间原样写入gmv_range或sales_volume_range；销售额指数写入gmv_index；"
                            "销量指数写入sales_index。指数绝不能写入gmv、sales_volume或order_count，"
                            "也绝不能用于求和、平均、客单价或GMV占比。只有页面明确标注为实际金额/件数"
                            "而不是指数时，才可写入gmv或sales_volume并把对应is_actual设为true。"
                            "sessions和product_breakdown必须逐条保留页面中能验证的记录。"
                            "报告需覆盖人设定位、内容策略、带货策略、粉丝画像、近期异动、"
                            "亮点短板、合作潜力和同行可复用打法。"
                        ),
                    },
                    {
                        "role": "user",
                        "content": "JSON结构示例："
                        + json.dumps(schema, ensure_ascii=False)
                        + "\n输入证据："
                        + json.dumps(prompt, ensure_ascii=False, default=str),
                    },
                ],
            },
        )
        parsed = _clean_json(response["choices"][0]["message"]["content"])
        parsed["_numeric_source"] = "llm"
        merged = {**fallback, **{key: value for key, value in parsed.items() if key in schema}}
        merged["_numeric_source"] = "llm"
        finalized = finalize_commerce_analysis(candidate, merged)
        finalized["web_snapshot"] = _evidence_payload(evidence)
        return finalized
    except Exception as exc:
        fallback["data_notes"].append(f"AI分析暂不可用，已使用规则版报告：{exc}")
        return fallback


def collect_chanmama_evidence(candidate: dict[str, Any], settings: Settings) -> dict[str, Any]:
    analysis_url = str(candidate.get("analysis_url") or "").strip()
    if not analysis_url:
        return {}
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        raise RuntimeError("缺少Playwright，无法读取蝉妈妈达人详情页") from exc
    url = urljoin("https://www.chanmama.com/", analysis_url)
    with sync_playwright() as playwright:
        attached_browser = None
        context = None
        owns_context = False
        page = None
        try:
            # The dedicated login Chrome normally remains open and owns the
            # persistent profile.  Reuse it over its loopback CDP endpoint;
            # starting a second Chrome with the same profile makes Chrome
            # terminate immediately and used to produce an unreadable wall of
            # ``Target page, context or browser has been closed`` logs.
            try:
                attached_browser = playwright.chromium.connect_over_cdp(
                    f"http://127.0.0.1:{CHANMAMA_CDP_PORT}", timeout=10_000
                )
                if attached_browser.contexts:
                    context = attached_browser.contexts[0]
            except Exception:
                attached_browser = None
            if context is None:
                try:
                    context = playwright.chromium.launch_persistent_context(
                        str(settings.chanmama_profile_dir),
                        channel="chrome",
                        headless=True,
                        no_viewport=True,
                        args=["--disable-background-timer-throttling"],
                    )
                    owns_context = True
                except Exception as exc:
                    if "Target page, context or browser has been closed" in str(exc):
                        raise RuntimeError(
                            "无法连接蝉妈妈专用Chrome。请保持专用Chrome窗口打开，"
                            "回到“登录与连接”重新打开一次，再生成达人拆解。"
                        ) from exc
                    raise
            # Use a separate tab so collecting a report does not replace the
            # page the user is currently viewing in the dedicated browser.
            page = context.new_page()
            page.goto(url, wait_until="domcontentloaded", timeout=60_000)
            page.wait_for_timeout(4_000)
            initial_body = page.locator("body").inner_text(timeout=15_000)
            if (
                "register.html" in page.url
                or "当前用户未登录" in initial_body
                or ("注册/登录" in initial_body and "微信登录" in initial_body)
            ):
                raise RuntimeError("蝉妈妈登录状态已失效，请先在专用浏览器重新登录")
            snapshots: list[dict[str, Any]] = []
            warnings: list[str] = []
            seen: set[str] = set()

            def click_visible_text(label: str) -> bool:
                locator = page.get_by_text(label, exact=True)
                for index in range(min(locator.count(), 12)):
                    item = locator.nth(index)
                    if not item.is_visible():
                        continue
                    try:
                        item.click(timeout=5_000)
                        page.wait_for_timeout(1_200)
                        return True
                    except Exception:
                        continue
                return False

            def capture(section: str, tab_label: str) -> bool:
                try:
                    # 分段滚动触发当前栏目正常的懒加载；不调用网页隐藏接口。
                    for ratio in (0, 0.25, 0.5, 0.75, 1):
                        page.evaluate(
                            "(ratio) => window.scrollTo(0, "
                            "Math.max(0, (document.body.scrollHeight - innerHeight) * ratio))",
                            ratio,
                        )
                        page.wait_for_timeout(220)
                    page.evaluate("window.scrollTo(0, 0)")
                    payload = page.evaluate(
                        """() => {
                            const visible = element => {
                                const style = window.getComputedStyle(element);
                                const rect = element.getBoundingClientRect();
                                return style.display !== 'none' && style.visibility !== 'hidden'
                                    && rect.width > 0 && rect.height > 0;
                            };
                            const clean = value => (value || '').replace(/\\u00a0/g, ' ')
                                .replace(/[ \\t]+/g, ' ').trim();
                            const bodyText = clean(document.body.innerText || '');
                            const lines = [];
                            const lineSeen = new Set();
                            for (const line of bodyText.split(/\\n+/)) {
                                const text = clean(line);
                                if (!text || lineSeen.has(text)) continue;
                                lineSeen.add(text);
                                lines.push(text);
                            }
                            const tables = [];
                            const tableSeen = new Set();
                            for (const table of document.querySelectorAll('table')) {
                                if (!visible(table)) continue;
                                const headers = Array.from(table.querySelectorAll('thead th'))
                                    .map(cell => clean(cell.innerText)).filter(Boolean);
                                const rows = Array.from(table.querySelectorAll('tbody tr')).map(row =>
                                    Array.from(row.querySelectorAll('td,th'))
                                        .map(cell => clean(cell.innerText))
                                ).filter(row => row.some(Boolean));
                                if (!headers.length && !rows.length) continue;
                                const signature = JSON.stringify([headers, rows.slice(0, 3)]);
                                if (tableSeen.has(signature)) continue;
                                tableSeen.add(signature);
                                tables.push({ headers, rows });
                            }
                            const chartLabels = [];
                            const chartSeen = new Set();
                            for (const node of document.querySelectorAll('canvas, svg')) {
                                if (!visible(node)) continue;
                                const parent = node.parentElement;
                                const text = clean(parent ? parent.innerText : '');
                                if (!text || chartSeen.has(text)) continue;
                                chartSeen.add(text);
                                chartLabels.push(text.slice(0, 2000));
                            }
                            const active = Array.from(document.querySelectorAll(
                                '[aria-selected="true"], .is-active, .active, .selected'
                            )).filter(visible).map(node => clean(node.innerText))
                              .filter(text => text && text.length <= 80).slice(0, 40);
                            return { lines, tables, chart_labels: chartLabels, active_labels: active };
                        }"""
                    )
                    fingerprint = json.dumps(
                        [section, tab_label, payload.get("lines", [])[:80]],
                        ensure_ascii=False,
                    )
                    if fingerprint in seen:
                        return False
                    seen.add(fingerprint)
                    lines = payload.get("lines") or []
                    access_note = ""
                    access_terms = ("暂无权限", "权限不足", "开通会员", "升级会员", "暂无数据")
                    matched = [line for line in lines if any(term in line for term in access_terms)]
                    if matched:
                        access_note = "；".join(matched[:5])
                    snapshots.append(
                        {
                            "section": section,
                            "tab": tab_label,
                            "title": page.title(),
                            "url": page.url,
                            "captured_at": datetime.now().astimezone().isoformat(timespec="seconds"),
                            "period": "30天" if "30天" in payload.get("active_labels", []) else "",
                            "active_labels": payload.get("active_labels") or [],
                            "lines": lines,
                            "tables": payload.get("tables") or [],
                            "chart_labels": payload.get("chart_labels") or [],
                            "access_note": access_note,
                        }
                    )
                    return True
                except Exception as exc:
                    warnings.append(f"{section}/{tab_label}采集失败：{exc}")
                    return False

            def capture_with_pagination(section: str, tab_label: str) -> None:
                capture(section, tab_label)
                # 页面表格若有“下一页”，逐页读取当前账号权限下可见的全部行。
                for page_number in range(2, 301):
                    next_locators = page.locator(
                        ".el-pagination .btn-next:not([disabled]), "
                        "li.ant-pagination-next:not(.ant-pagination-disabled), "
                        "li.ivu-page-next:not(.ivu-page-disabled)"
                    )
                    next_button = None
                    for index in range(min(next_locators.count(), 12)):
                        item = next_locators.nth(index)
                        if item.is_visible() and item.is_enabled():
                            next_button = item
                            break
                    if next_button is None:
                        break
                    try:
                        next_button.click(timeout=5_000)
                        page.wait_for_timeout(1_000)
                    except Exception as exc:
                        warnings.append(f"{section}/{tab_label}翻页失败：{exc}")
                        break
                    if not capture(section, f"{tab_label} / 分页{page_number}"):
                        break

            def discover_tabs() -> list[str]:
                try:
                    values = page.locator(
                        "[role='tab'], .el-tabs__item, .ant-tabs-tab, .ivu-tabs-tab"
                    ).evaluate_all(
                        """elements => Array.from(new Set(elements
                            .filter(element => {
                                const style = window.getComputedStyle(element);
                                const rect = element.getBoundingClientRect();
                                return style.display !== 'none' && style.visibility !== 'hidden'
                                    && rect.width > 0 && rect.height > 0;
                            })
                            .map(element => (element.innerText || '').replace(/\\s+/g, ' ').trim())
                            .filter(text => text && text.length <= 24)))"""
                    )
                    return [str(value) for value in values]
                except Exception as exc:
                    warnings.append(f"子标签识别失败：{exc}")
                    return []

            # 优先固定为30天，确保栏目之间采用同一统计周期。
            click_visible_text("30天")
            click_visible_text("展开")
            capture_with_pagination("账号概览", "达人资料与榜单标签")

            # 只遍历数据分析侧栏与其标签；收藏、监控、建联等写操作不会触发。
            ignored_tabs = {
                "7天", "30天", "日", "周", "月", "自然", "大促",
                "全部", "刷新", "导出数据", "账号拆解",
            }
            captured_sections: list[str] = []
            for section in DETAIL_SECTIONS:
                if not click_visible_text(section):
                    warnings.append(f"未找到或无法打开栏目：{section}")
                    continue
                captured_sections.append(section)
                click_visible_text("30天")
                capture_with_pagination(section, "栏目首页")
                tab_labels = list(
                    dict.fromkeys(
                        [*SECTION_TABS.get(section, ()), *discover_tabs()]
                    )
                )
                for tab_label in tab_labels[:60]:
                    if tab_label in ignored_tabs or tab_label == section:
                        continue
                    if click_visible_text(tab_label):
                        capture_with_pagination(section, tab_label)
                        try:
                            page.keyboard.press("Escape")
                        except Exception:
                            pass
                    elif tab_label in SECTION_TABS.get(section, ()):
                        warnings.append(f"{section}未找到或无法打开页签：{tab_label}")

            return {
                "source": "蝉妈妈达人详情页（当前账号权限下可见页面）",
                "captured_at": datetime.now().astimezone().isoformat(timespec="seconds"),
                "analysis_url": url,
                "coverage": {
                    "attempted_sections": list(DETAIL_SECTIONS),
                    "captured_sections": list(dict.fromkeys(captured_sections)),
                    "missing_sections": [
                        section for section in DETAIL_SECTIONS if section not in captured_sections
                    ],
                    "snapshot_count": len(snapshots),
                },
                "leaderboard": {
                    "raw_data": _raw_candidate_data(candidate),
                    "rank": _raw_candidate_data(candidate).get("排行"),
                    "anchor_name": candidate.get("anchor_name"),
                    "douyin_id": candidate.get("douyin_id"),
                    "gmv_range": candidate.get("estimated_gmv_text"),
                    "gmv_index": candidate.get("gmv_index"),
                    "sales_range": candidate.get("sales_volume_text"),
                    "sales_index": candidate.get("sales_index"),
                    "sales_aov": _candidate_aov(candidate),
                    "followers": candidate.get("followers"),
                    "category": candidate.get("category"),
                    "sessions": candidate.get("sessions_7d"),
                },
                "pages": snapshots,
                "warnings": list(dict.fromkeys(warnings)),
            }
        finally:
            if page is not None:
                try:
                    page.close()
                except Exception:
                    pass
            # Never close a context that belongs to the user's dedicated
            # Chrome.  Only shut down the private fallback context we created.
            if owns_context and context is not None:
                context.close()


def _set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell, top: int = 120, start: int = 140, bottom: int = 120, end: int = 140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size in (("Heading 1", 15), ("Heading 2", 12)):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(13)
        style.paragraph_format.space_after = Pt(6)


def _add_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Heading 1"]
    marker = paragraph.add_run("▌")
    marker.font.color.rgb = RGBColor.from_string(ACCENT)
    marker.bold = True
    paragraph.add_run(title)


def _add_bullets(document: Document, items: list[Any]) -> None:
    values = [str(item).strip() for item in (items or []) if str(item).strip()]
    if not values:
        values = ["数据不足"]
    for item in values:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def _add_metrics(document: Document, metrics: list[list[Any]]) -> None:
    rows = [row for row in (metrics or []) if isinstance(row, list) and len(row) >= 2]
    if not rows:
        rows = [["数据状态", "数据不足"]]
    columns = 2
    table = document.add_table(rows=(len(rows) + 1) // 2, cols=columns * 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.1), Cm(5.2), Cm(3.1), Cm(5.2)]
    for row_index, table_row in enumerate(table.rows):
        for cell_index, cell in enumerate(table_row.cells):
            cell.width = widths[cell_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_fill(cell, LIGHT)
            _set_cell_margins(cell)
            source_index = row_index * 2 + cell_index // 2
            if source_index >= len(rows):
                cell.text = ""
                continue
            value = str(rows[source_index][cell_index % 2])
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            if cell_index % 2 == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(MUTED)
            else:
                run.bold = True
                run.font.size = Pt(12)


def _add_comparison(document: Document, strengths: list[Any], risks: list[Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = [("优势与亮点", GREEN), ("短板与风险", AMBER)]
    for index, (label, color) in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Cm(8.35)
        _set_cell_fill(cell, LIGHT)
        _set_cell_margins(cell, 160, 180, 160, 180)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(color)
        items = strengths if index == 0 else risks
        for item in items or ["数据不足"]:
            paragraph = cell.add_paragraph(style="List Bullet")
            paragraph.add_run(str(item))


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        properties.append(marker)
    marker.set(qn("w:val"), "true")


def _add_hyperlink(paragraph, url: str, label: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_data_table(
    document: Document,
    headers: list[str],
    rows: list[list[Any]],
    widths_cm: list[float],
    *,
    link_columns: set[int] | None = None,
) -> None:
    link_columns = link_columns or set()
    if not rows:
        rows = [["数据不足", *(["—"] * (len(headers) - 1))]]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table._tbl.tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(Cm(width).twips))
        grid.append(grid_col)
    _set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Cm(widths_cm[index])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_fill(cell, ACCENT)
        _set_cell_margins(cell, 110, 100, 110, 100)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row_values in rows:
        row = table.add_row()
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell, 105, 100, 105, 100)
            value = row_values[index] if index < len(row_values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index in {0, len(headers) - 1} else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            if index in link_columns and isinstance(value, str) and value.startswith(("http://", "https://")):
                _add_hyperlink(paragraph, value, "打开链接")
            else:
                run = paragraph.add_run(str(value))
                run.font.size = Pt(8.2)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _add_leaderboard_snapshot(
    document: Document,
    candidate: dict[str, Any],
    analysis: dict[str, Any],
) -> None:
    snapshot = _evidence_payload(analysis.get("web_snapshot"))
    leaderboard = snapshot.get("leaderboard") if isinstance(snapshot.get("leaderboard"), dict) else {}
    raw = _raw_candidate_data(candidate)
    rows: list[list[Any]] = [
        ["榜单排名", raw.get("排行") or leaderboard.get("rank") or "数据不足"],
        ["达人", candidate.get("anchor_name") or "数据不足"],
        ["抖音号", candidate.get("douyin_id") or "数据不足"],
        [
            "直播销售额（页面区间）",
            _format_money_metric(candidate.get("estimated_gmv"), candidate.get("estimated_gmv_text")),
        ],
        ["销售额指数（仅供比较）", _format_index(candidate.get("gmv_index"))],
        [
            "直播销量（页面区间）",
            _format_quantity(candidate.get("sales_volume"), candidate.get("sales_volume_text")),
        ],
        ["销量指数（仅供比较）", _format_index(candidate.get("sales_index"))],
        ["销售客单价", _candidate_aov(candidate)],
        ["粉丝数", format_metric(candidate.get("followers"))],
        ["带货类目", candidate.get("category") or "数据不足"],
        ["直播场次", format_metric(candidate.get("sessions_7d"))],
        ["直播状态", raw.get("直播间") or raw.get("直播状态") or "数据不足"],
    ]
    represented = {
        "排行", "达人", "抖音号", "直播销售额(元)", "销售额指数",
        "直播销量(件)", "销量指数", "销售客单价", "粉丝数",
        "带货类目", "直播场次", "直播间", "直播状态",
    }
    for key, value in raw.items():
        if str(key) in represented or value in (None, ""):
            continue
        rows.append([str(key), str(value)])
    document.add_page_break()
    _add_heading(document, "榜单原始数据（完整字段）")
    _add_bullets(
        document,
        [
            "页面大字区间属于成交数据；灰色小字指数仅用于同区间比较，报告不会把指数当作实际金额或件数。",
            f"榜单数据来源：{candidate.get('source') or '蝉妈妈'}。",
        ],
    )
    _add_data_table(document, ["字段", "页面原始值"], rows, [5.0, 11.5])


def _add_subheading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Heading 2"]
    paragraph.add_run(title)


def _clean_appendix_lines(lines: Any) -> list[str]:
    if not isinstance(lines, list):
        return []
    ignored = {
        "首页", "达人", "商品", "直播", "视频", "小店", "品牌", "工具",
        "蝉妈妈创意", "APP", "企业版", "添加对比", "建联", "找相似",
        "监控", "账号拆解", "刷新", "展开",
    }
    result: list[str] = []
    seen: set[str] = set()
    for value in lines:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text or text in ignored or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return result


def _add_full_web_appendix(document: Document, analysis: dict[str, Any]) -> None:
    snapshot = _evidence_payload(analysis.get("web_snapshot"))
    pages = [page for page in (snapshot.get("pages") or []) if isinstance(page, dict)]
    coverage = snapshot.get("coverage") if isinstance(snapshot.get("coverage"), dict) else {}
    if not snapshot:
        return
    landscape = document.add_section(WD_SECTION.NEW_PAGE)
    _set_landscape(landscape)
    _add_heading(document, "蝉妈妈账号详情完整数据附录")
    _add_bullets(
        document,
        [
            f"采集时间：{snapshot.get('captured_at') or '数据不足'}。",
            f"详情页：{snapshot.get('analysis_url') or '数据不足'}。",
            f"已采集页面/标签：{len(pages)} 个；仅包含当前账号权限下网页正常展示的数据。",
            "已进入栏目："
            + "、".join(str(item) for item in (coverage.get("captured_sections") or []))
            if coverage.get("captured_sections")
            else "已进入栏目：以附录逐项记录为准。",
            "附录保留页面原始区间、指数、文本、表格及图表标签；未展示的数据不会推测。",
        ],
    )
    warnings = [str(item) for item in (snapshot.get("warnings") or []) if str(item).strip()]
    if warnings:
        _add_subheading(document, "采集警告与权限边界")
        _add_bullets(document, warnings)

    seen_table_signatures: set[str] = set()
    seen_page_lines: set[str] = set()
    for page_index, page in enumerate(pages, 1):
        section = str(page.get("section") or "未命名栏目")
        tab_label = str(page.get("tab") or "栏目首页")
        _add_subheading(document, f"{page_index}. {section} / {tab_label}")
        metadata = [
            f"统计周期：{page.get('period') or '以页面当前选中周期为准'}",
            f"页面标题：{page.get('title') or '数据不足'}",
            f"页面地址：{page.get('url') or '数据不足'}",
            f"采集时间：{page.get('captured_at') or '数据不足'}",
        ]
        if page.get("access_note"):
            metadata.append(f"访问说明：{page['access_note']}")
        _add_bullets(document, metadata)

        tables = [table for table in (page.get("tables") or []) if isinstance(table, dict)]
        for table_index, table in enumerate(tables, 1):
            source_rows = [row for row in (table.get("rows") or []) if isinstance(row, list)]
            headers = [str(value) for value in (table.get("headers") or [])]
            table_signature = json.dumps(
                [headers, source_rows],
                ensure_ascii=False,
                default=str,
            )
            if table_signature in seen_table_signatures:
                continue
            seen_table_signatures.add(table_signature)
            column_count = max(
                [len(headers), *(len(row) for row in source_rows)] if source_rows else [len(headers)]
            )
            if column_count <= 0:
                continue
            if not headers:
                headers = [f"列{index + 1}" for index in range(column_count)]
            elif len(headers) < column_count:
                headers.extend(f"列{index + 1}" for index in range(len(headers), column_count))
            normalized_rows = [
                [str(row[index]) if index < len(row) else "" for index in range(column_count)]
                for row in source_rows
            ]
            widths = [25.3 / column_count] * column_count
            _add_subheading(document, f"数据表 {table_index}")
            _add_data_table(document, headers, normalized_rows, widths)

        lines = _clean_appendix_lines(page.get("lines"))
        chart_labels = _clean_appendix_lines(page.get("chart_labels"))
        for label in chart_labels:
            if label not in lines:
                lines.append("图表标签：" + label)
        unique_lines: list[str] = []
        for line in lines:
            if line in seen_page_lines:
                continue
            seen_page_lines.add(line)
            unique_lines.append(line)
        lines = unique_lines
        if lines:
            _add_subheading(document, "页面字段与文字原文")
            _add_data_table(
                document,
                ["序号", "页面原文"],
                [[index, line] for index, line in enumerate(lines, 1)],
                [1.5, 23.8],
            )

def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def _set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def _add_session_and_product_sections(document: Document, analysis: dict[str, Any]) -> None:
    landscape = document.add_section(WD_SECTION.NEW_PAGE)
    _set_landscape(landscape)

    period = analysis.get("period_summary") or {}
    stability = period.get("stability") or {}
    _add_heading(document, "近30天/当前可见周期场次表现")
    _add_bullets(
        document,
        [
            f"统计周期：{period.get('period') or '数据不足'}；周期GMV："
            f"{_format_money_metric(period.get('total_gmv'), period.get('total_gmv_text'))}。",
            (
                f"取得逐场GMV {stability.get('gmv_session_count', 0)} 场；"
                f"场均GMV {_format_currency(stability.get('average_gmv'))}；"
                f"GMV中位数 {_format_currency(stability.get('median_gmv'))}。"
            ),
            (
                f"最高单场GMV {_format_currency(stability.get('max_gmv'))}，"
                f"占已取得逐场GMV的 {_format_percent(stability.get('max_share'))}；"
                f"判断：{stability.get('label') or '数据不足'}。"
            ),
        ],
    )
    session_rows = [
        [
            row.get("date") or "时间未取得",
            _format_money_metric(row.get("gmv"), row.get("gmv_range")),
            _format_quantity(row.get("sales_volume"), row.get("sales_volume_range")),
            format_metric(row.get("order_count")),
            _format_price(row.get("aov")),
            row.get("duration") or "数据不足",
            row.get("products") or "数据不足",
            row.get("source") or "数据不足",
        ]
        for row in analysis.get("sessions", [])
    ]
    _add_data_table(
        document,
        ["日期/时间", "场次GMV", "销量（件）", "订单", "场次客单价", "时长", "本场主要商品", "数据来源"],
        session_rows,
        [2.8, 2.3, 1.7, 1.7, 2.2, 1.8, 8.0, 4.0],
    )

    _add_heading(document, "商品结构拆解")
    product_rows = [
        [
            row.get("product_name") or "商品名未取得",
            row.get("category") or "数据不足",
            row.get("product_url") or "数据不足",
            _format_price(row.get("aov")),
            _format_money_metric(row.get("gmv"), row.get("gmv_range")),
            _format_percent(row.get("gmv_share")),
            _format_quantity(row.get("sales_volume"), row.get("sales_volume_range")),
            format_metric(row.get("session_count")),
            row.get("source") or "数据不足",
        ]
        for row in analysis.get("product_breakdown", [])
    ]
    _add_data_table(
        document,
        ["商品", "细分品类", "商品链接", "成交均价", "商品GMV", "GMV占比", "销量", "出现场次", "数据来源"],
        product_rows,
        [5.0, 2.7, 2.2, 2.2, 2.3, 2.0, 1.8, 1.8, 4.0],
        link_columns={2},
    )

    portrait = document.add_section(WD_SECTION.NEW_PAGE)
    _set_portrait(portrait)


def build_report_docx(
    candidate: dict[str, Any],
    analysis: dict[str, Any],
    output_path: Path,
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(5)
    anchor_name = str(candidate.get("anchor_name") or "达人")
    report_subject = anchor_name if anchor_name.endswith("达人") else anchor_name + "达人"
    run = title.add_run(f"{report_subject}深度分析报告")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(INK)
    subtitle = document.add_paragraph(str(analysis.get("subtitle") or "直播电商达人拆解"))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    meta = document.add_paragraph(
        f"生成时间：{datetime.now().astimezone().strftime('%Y-%m-%d %H:%M')}  "
        f"｜数据来源：{candidate.get('source') or '蝉妈妈'}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(8.5)
    meta.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    _add_heading(document, "达人基础定位")
    _add_bullets(document, analysis.get("positioning", []))
    _add_heading(document, "核心数据概览")
    _add_metrics(document, analysis.get("metrics", []))
    _add_leaderboard_snapshot(document, candidate, analysis)
    _add_session_and_product_sections(document, analysis)
    _add_heading(document, "人设定位")
    _add_bullets(document, analysis.get("persona", []))
    _add_heading(document, "内容策略")
    _add_bullets(document, analysis.get("content_strategy", []))
    _add_heading(document, "带货与商品策略")
    _add_bullets(document, analysis.get("commerce_strategy", []))
    _add_heading(document, "粉丝受众画像")
    _add_bullets(document, analysis.get("audience", []))
    _add_heading(document, "近期异动与关键发现")
    _add_bullets(document, analysis.get("recent_changes", []))
    _add_heading(document, "优势与短板分析")
    _add_comparison(document, analysis.get("strengths", []), analysis.get("risks", []))
    _add_heading(document, "达人合作潜力")
    _add_bullets(document, analysis.get("cooperation", []))
    _add_heading(document, "同行可复用的关键打法")
    _add_bullets(document, analysis.get("reusable_playbook", []))
    _add_heading(document, "数据完整性说明")
    _add_bullets(document, analysis.get("data_notes", []))

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not footer.text:
            footer_run = footer.add_run("直播主播发现与拆解 Agent")
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor.from_string(MUTED)
    document.save(output_path)
    return output_path


class CreatorReportManager:
    def __init__(self, settings: Settings, database: Database):
        self.settings = settings
        self.database = database
        # start() returns status() while holding this lock, so it must be
        # re-entrant. A plain Lock deadlocks the request before the UI receives
        # the "report started" response.
        self._lock = threading.RLock()
        self._thread: threading.Thread | None = None
        self._state: dict[str, Any] = {
            "busy": False,
            "phase": "idle",
            "message": "尚未生成达人拆解报告",
            "reports": [],
        }

    def status(self) -> dict[str, Any]:
        with self._lock:
            return json.loads(json.dumps(self._state, ensure_ascii=False))

    def start(self, candidate_ids: list[int]) -> dict[str, Any]:
        with self._lock:
            if self._thread is not None and self._thread.is_alive():
                raise RuntimeError("达人拆解任务正在运行，请等待完成")
            all_candidates = self.database.list_candidates(limit=5000)
            candidates = [item for item in all_candidates if int(item["id"]) in candidate_ids]
            if not candidates:
                raise RuntimeError("请先选择需要拆解的达人")
            self._state = {
                "busy": True,
                "phase": "collecting",
                "message": f"正在准备拆解 {len(candidates)} 位达人",
                "reports": [],
                "total": len(candidates),
                "completed": 0,
            }
            self._thread = threading.Thread(target=self._run, args=(candidates,), daemon=True)
            self._thread.start()
            return self.status()

    def _update(self, **values: Any) -> None:
        with self._lock:
            self._state.update(values)

    def _run(self, candidates: list[dict[str, Any]]) -> None:
        reports: list[dict[str, Any]] = []
        try:
            for index, candidate in enumerate(candidates, 1):
                name = str(candidate.get("anchor_name") or f"达人{candidate['id']}")
                self._update(
                    phase="collecting",
                    message=f"正在读取 {name} 的蝉妈妈详情数据（{index}/{len(candidates)}）",
                )
                evidence = validate_detail_evidence(
                    collect_chanmama_evidence(candidate, self.settings)
                )
                evidence_name = f"{safe_name(name)}_蝉妈妈网页完整数据.json"
                evidence_path = self.settings.report_dir / evidence_name
                evidence_temp_path = self.settings.report_dir / (
                    f".{safe_name(name)}_蝉妈妈网页完整数据.tmp.json"
                )
                evidence_temp_path.write_text(
                    json.dumps(evidence, ensure_ascii=False, indent=2, default=str),
                    encoding="utf-8",
                )
                evidence_temp_path.replace(evidence_path)
                self._update(
                    phase="analyzing",
                    message=f"正在分析 {name} 的逐场GMV、商品结构与带货打法",
                )
                analysis = analyze_creator(candidate, evidence, self.settings)
                file_name = f"{safe_name(name)}_达人拆解报告.docx"
                path = self.settings.report_dir / file_name
                report_temp_path = self.settings.report_dir / (
                    f".{safe_name(name)}_达人拆解报告.tmp.docx"
                )
                build_report_docx(candidate, analysis, report_temp_path)
                report_temp_path.replace(path)
                reports.append(
                    {
                        "candidate_id": candidate["id"],
                        "anchor_name": name,
                        "file_name": path.name,
                        "download_url": "/reports/" + quote(path.name),
                        "evidence_file_name": evidence_path.name,
                        "evidence_download_url": "/reports/" + quote(evidence_path.name),
                        "used_detail_page": bool(evidence),
                    }
                )
                self.database.update_candidate_status([int(candidate["id"])], "analyzed")
                self._update(reports=reports, completed=len(reports))
            self._update(
                busy=False,
                phase="completed",
                message=f"已生成 {len(reports)} 份达人拆解报告",
                reports=reports,
            )
        except Exception as exc:
            self._update(busy=False, phase="error", message=f"达人拆解失败：{exc}", reports=reports)
